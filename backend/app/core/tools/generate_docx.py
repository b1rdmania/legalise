"""generate_docx tool — render a markdown body to a .docx file.

v0.1 markdown handling is intentionally narrow:
  - blank lines separate paragraphs
  - `# ` / `## ` / `### ` prefixes become Heading 1/2/3
  - single newlines inside a paragraph become soft line breaks

A full markdown parser is out of scope; the tool's job is to produce a
plausibly-formatted Word document, not perfect fidelity.

Unit 1: bytes are written to object storage via `app.core.storage`.
The `storage_uri` returned and stored in audit payloads is the object key
(``users/{user_id}/matters/{matter_id}/generated/{file_uuid}/{file_uuid}.docx``
or a legacy key when user_id is not available).
The download endpoint in documents.py reads from storage, not the filesystem.
"""

from __future__ import annotations

import io
import uuid

from docx import Document as DocxDocument
from docx.enum.section import WD_ORIENTATION
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.api import audit, audit_failure
from app.core.storage import get_storage_backend, generated_key, StorageWriteError
from app.core.tools.schemas import GenerateDocxInput, GenerateDocxOutput


def _apply_orientation(document: DocxDocument, orientation: str) -> None:
    if orientation != "landscape":
        return
    section = document.sections[0]
    new_width, new_height = section.page_height, section.page_width
    section.orientation = WD_ORIENTATION.LANDSCAPE
    section.page_width = new_width
    section.page_height = new_height


def _render_markdown(document: DocxDocument, title: str, body: str) -> None:
    document.add_heading(title, level=0)
    for block in body.split("\n\n"):
        block = block.rstrip()
        if not block:
            continue
        if block.startswith("### "):
            document.add_heading(block[4:].strip(), level=3)
        elif block.startswith("## "):
            document.add_heading(block[3:].strip(), level=2)
        elif block.startswith("# "):
            document.add_heading(block[2:].strip(), level=1)
        else:
            lines = block.split("\n")
            para = document.add_paragraph(lines[0])
            for line in lines[1:]:
                run = para.add_run()
                run.add_break()
                para.add_run(line)


async def handle_generate_docx(
    inputs: GenerateDocxInput,
    *,
    session: AsyncSession,
    actor_id: uuid.UUID,
    matter_id: uuid.UUID | None,
) -> GenerateDocxOutput:
    document = DocxDocument()
    orientation = (inputs.options.orientation if inputs.options else "portrait")
    _apply_orientation(document, orientation)
    _render_markdown(document, inputs.title, inputs.body_markdown)

    # Serialise to bytes in memory — no filesystem dependency.
    buf = io.BytesIO()
    document.save(buf)
    docx_bytes = buf.getvalue()
    byte_count = len(docx_bytes)
    char_count = len(inputs.body_markdown)

    # Resolve matter_id from options if not provided directly.
    effective_matter_id = matter_id
    if effective_matter_id is None and inputs.options and inputs.options.matter_id:
        effective_matter_id = inputs.options.matter_id

    file_uuid = uuid.uuid4()
    filename = f"{file_uuid}.docx"

    # Build storage key. Unit 1: actor_id + matter_id are both available
    # here so we use the canonical generated_key shape. When matter_id is
    # genuinely absent (_orphan case) we fall back to a legacy key so the
    # tool does not break callers that don't supply it.
    if effective_matter_id is not None:
        key = generated_key(
            user_id=actor_id,
            matter_id=effective_matter_id,
            document_id=file_uuid,
            filename=filename,
        )
    else:
        key = f"generated/_orphan/{file_uuid}/{filename}"

    storage = get_storage_backend()
    try:
        storage.put_bytes(
            key,
            docx_bytes,
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            metadata={"title": inputs.title[:200], "orientation": orientation},
        )
    except StorageWriteError as exc:
        # Failure-path audit row via `audit_failure` (separate committed
        # session). The caller route raises HTTPException without
        # committing, so a `session.add` here would be lost at teardown.
        # R3 round-2 review fix — same pattern as the upload path in
        # matters.py and the download path in documents.py.
        await audit_failure(
            session,
            "storage.put_bytes.failed",
            actor_id=actor_id,
            matter_id=effective_matter_id,
            module="storage",
            resource_type="document",
            resource_id=str(file_uuid),
            payload={
                "storage_key": key,
                "backend": exc.backend,
                "error_code": exc.error_code,
            },
        )
        raise

    storage_uri = key

    await audit.log(
        session,
        "document.generated",
        actor_id=actor_id,
        matter_id=effective_matter_id,
        module="document_generation",
        resource_type="document",
        resource_id=str(file_uuid),
        payload={
            "format": "docx",
            "char_count": char_count,
            "byte_count": byte_count,
            "storage_uri": storage_uri,
            "orientation": orientation,
            "title": inputs.title,
        },
    )

    return GenerateDocxOutput(
        storage_uri=storage_uri,
        byte_count=byte_count,
        char_count=char_count,
    )


__all__ = ["handle_generate_docx"]
