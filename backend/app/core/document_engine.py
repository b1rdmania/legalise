"""Document engine primitives.

Legalise owns the document workspace contract rather than vendoring a full
editor backend. Skills and UI code ask for one structured snapshot of a
document, then render or operate on that snapshot.
"""

from __future__ import annotations

import io
import uuid
from dataclasses import dataclass, field
from typing import Literal

from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.storage import StorageReadError, get_storage_backend
from app.models import Document, DocumentVersion, Matter, STATUS_ARCHIVED
from app.models.document_body import extracted_body_for


BlockType = Literal["paragraph", "table_cell"]
SnapshotSource = Literal["original_docx", "latest_version", "extracted_body"]


@dataclass(frozen=True)
class DocumentBlock:
    id: str
    type: BlockType
    ordinal: int
    text: str


@dataclass(frozen=True)
class DocumentSnapshot:
    document: Document
    matter: Matter
    source: SnapshotSource
    source_version: DocumentVersion | None
    extraction_method: str | None
    blocks: list[DocumentBlock]
    text: str
    notes: list[str] = field(default_factory=list)

    @property
    def char_count(self) -> int:
        return len(self.text)


class DocumentEngineNotFound(LookupError):
    """Document is absent, archived, or not owned by the actor."""


class DocumentEngineUnavailable(ValueError):
    """Document exists but has no readable text yet."""


def _normalise_blocks(blocks: list[DocumentBlock]) -> list[DocumentBlock]:
    return [block for block in blocks if block.text.strip()]


def blocks_from_text(text: str) -> list[DocumentBlock]:
    """Split readable text into stable paragraph blocks."""
    parts = [part.strip() for part in text.replace("\r\n", "\n").split("\n\n")]
    blocks: list[DocumentBlock] = []
    for index, part in enumerate(part for part in parts if part):
        blocks.append(
            DocumentBlock(
                id=f"p{index + 1}",
                type="paragraph",
                ordinal=index + 1,
                text=part,
            )
        )
    return blocks


def blocks_from_docx(file_bytes: bytes) -> list[DocumentBlock]:
    """Extract paragraphs and table cells from a DOCX file."""
    from docx import Document as DocxDocument

    document = DocxDocument(io.BytesIO(file_bytes))
    blocks: list[DocumentBlock] = []
    ordinal = 1

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            blocks.append(
                DocumentBlock(
                    id=f"p{ordinal}",
                    type="paragraph",
                    ordinal=ordinal,
                    text=text,
                )
            )
            ordinal += 1

    for table_index, table in enumerate(document.tables, start=1):
        for row_index, row in enumerate(table.rows, start=1):
            for cell_index, cell in enumerate(row.cells, start=1):
                text = cell.text.strip()
                if not text:
                    continue
                blocks.append(
                    DocumentBlock(
                        id=f"t{table_index}r{row_index}c{cell_index}",
                        type="table_cell",
                        ordinal=ordinal,
                        text=text,
                    )
                )
                ordinal += 1

    return _normalise_blocks(blocks)


async def _latest_resolved_version(
    session: AsyncSession, document_id: uuid.UUID
) -> DocumentVersion | None:
    return await session.scalar(
        select(DocumentVersion)
        .where(
            DocumentVersion.document_id == document_id,
            DocumentVersion.resolved_text.is_not(None),
        )
        .order_by(
            DocumentVersion.version_number.desc(),
            DocumentVersion.created_at.desc(),
        )
        .limit(1)
    )


def _is_docx(document: Document) -> bool:
    return document.filename.lower().endswith(".docx") or document.mime_type == (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


async def load_document_snapshot(
    session: AsyncSession,
    *,
    document_id: uuid.UUID,
    actor_id: uuid.UUID,
) -> DocumentSnapshot:
    """Return the current structured document snapshot for an owner.

    Source precedence:
    1. Latest resolved version text. Accepted edits define the current text.
    2. Original DOCX blocks if no resolved version exists and object storage is
       available.
    3. Extracted body fallback.
    """
    pair = (
        await session.execute(
            select(Document, Matter)
            .join(Matter, Matter.id == Document.matter_id)
            .where(Document.id == document_id)
        )
    ).first()
    if pair is None:
        raise DocumentEngineNotFound("document not found")

    document, matter = pair
    if matter.created_by_id != actor_id or matter.status == STATUS_ARCHIVED:
        raise DocumentEngineNotFound("document not found")

    latest_version = await _latest_resolved_version(session, document_id)
    if latest_version and latest_version.resolved_text:
        blocks = blocks_from_text(latest_version.resolved_text)
        if not blocks:
            raise DocumentEngineUnavailable("document has no readable text")
        return DocumentSnapshot(
            document=document,
            matter=matter,
            source="latest_version",
            source_version=latest_version,
            extraction_method=None,
            blocks=blocks,
            text="\n\n".join(block.text for block in blocks),
        )

    notes: list[str] = []
    if _is_docx(document) and document.storage_uri:
        try:
            original_bytes = get_storage_backend().get_bytes(document.storage_uri)
            blocks = blocks_from_docx(original_bytes)
            if blocks:
                return DocumentSnapshot(
                    document=document,
                    matter=matter,
                    source="original_docx",
                    source_version=None,
                    extraction_method="python-docx",
                    blocks=blocks,
                    text="\n\n".join(block.text for block in blocks),
                )
            notes.append("original DOCX had no readable blocks")
        except (KeyError, StorageReadError):
            notes.append("original DOCX unavailable; using extracted body")
        except Exception:  # noqa: BLE001
            notes.append("original DOCX could not be parsed; using extracted body")

    body = await extracted_body_for(session, document_id)
    if body is None or body.extraction_method == "failed" or not body.extracted_text.strip():
        raise DocumentEngineUnavailable("document body not available")

    blocks = blocks_from_text(body.extracted_text)
    if not blocks:
        raise DocumentEngineUnavailable("document has no readable text")

    return DocumentSnapshot(
        document=document,
        matter=matter,
        source="extracted_body",
        source_version=None,
        extraction_method=body.extraction_method,
        blocks=blocks,
        text="\n\n".join(block.text for block in blocks),
        notes=notes,
    )
