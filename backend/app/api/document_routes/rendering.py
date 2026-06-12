"""Documents API — DOCX / PDF / HTML rendering and export-filename helpers."""

from __future__ import annotations

import base64
import html
import io
import re
import uuid
from datetime import UTC, datetime
from typing import Any
from urllib.parse import unquote, urlparse

import httpx
from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.shared import Inches, RGBColor

from app.core.config import settings
from app.core.storage import (
    StorageReadError,
    document_asset_key,
    get_storage_backend,
)
from app.models import COMMENT_STATUS_RESOLVED, DocumentComment, DocumentVersion

from .schemas import DOCUMENT_ASSET_URL_RE, DocumentAssetContext

_FILENAME_SAFE_RE = re.compile(r"[^A-Za-z0-9._-]+")


def _safe_filename(title: str | None, file_uuid: str) -> str:
    """Derive a human-readable .docx filename from the audit `title` payload.

    Falls back to `generated-{uuid8}.docx` when title is empty or sanitises
    to nothing. The slugified title is bounded at 80 chars.
    """
    if title:
        cleaned = _FILENAME_SAFE_RE.sub("-", title).strip("-._")
        cleaned = cleaned[:80].rstrip("-._")
        if cleaned:
            return f"{cleaned}.docx"
    return f"generated-{file_uuid[:8]}.docx"


def _safe_asset_filename(filename: str | None, fallback: str) -> str:
    if filename:
        cleaned = _FILENAME_SAFE_RE.sub("-", filename).strip("-._")
        cleaned = cleaned[:100].rstrip("-._")
        if cleaned:
            return cleaned
    return fallback


def _docx_export_filename(filename: str, version_number: int) -> str:
    stem = filename.rsplit(".", 1)[0] if "." in filename else filename
    cleaned = _FILENAME_SAFE_RE.sub("-", stem).strip("-._")[:80].rstrip("-._")
    if not cleaned:
        cleaned = "document"
    return f"{cleaned}-v{version_number}.docx"


def _pdf_export_filename(filename: str, version_number: int) -> str:
    stem = filename.rsplit(".", 1)[0] if "." in filename else filename
    cleaned = _FILENAME_SAFE_RE.sub("-", stem).strip("-._")[:80].rstrip("-._")
    if not cleaned:
        cleaned = "document"
    return f"{cleaned}-v{version_number}.pdf"


def _render_resolved_text_docx(
    title: str,
    body: str,
    comments: list[DocumentComment] | None = None,
) -> bytes:
    document = DocxDocument()
    document.add_heading(title, level=0)
    for block in body.split("\n\n"):
        block = block.rstrip()
        if not block:
            continue
        lines = block.split("\n")
        para = document.add_paragraph(lines[0])
        for line in lines[1:]:
            para.add_run().add_break()
            para.add_run(line)
    buf = io.BytesIO()
    _append_document_comments_docx(document, comments or [])
    document.save(buf)
    return buf.getvalue()


def _append_document_comments_docx(
    document: DocxDocument,
    comments: list[DocumentComment],
) -> None:
    if not comments:
        return
    document.add_page_break()
    document.add_heading("Document review notes", level=1)
    for index, comment in enumerate(comments, start=1):
        status = "resolved" if comment.status == COMMENT_STATUS_RESOLVED else "open"
        document.add_heading(f"Note {index} ({status})", level=2)
        if comment.quote_text:
            quote = document.add_paragraph(style="Intense Quote")
            quote.add_run(comment.quote_text)
        document.add_paragraph(comment.body)
        meta = document.add_paragraph()
        meta.add_run("Created: ").bold = True
        meta.add_run(comment.created_at.isoformat())
        if comment.resolved_at:
            meta.add_run(" · Resolved: ").bold = True
            meta.add_run(comment.resolved_at.isoformat())


def _rgb_from_hex(value: str | None) -> RGBColor | None:
    if not value:
        return None
    cleaned = value.strip().lstrip("#")
    if len(cleaned) != 6:
        return None
    try:
        return RGBColor(
            int(cleaned[0:2], 16),
            int(cleaned[2:4], 16),
            int(cleaned[4:6], 16),
        )
    except ValueError:
        return None


def _paragraph_alignment(value: str | None):
    return {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }.get(value or "")


def _document_asset_key_from_src(
    src: str | None,
    context: DocumentAssetContext | None,
) -> str | None:
    if not src or context is None:
        return None
    parsed = urlparse(src)
    path = parsed.path if parsed.scheme or parsed.netloc else src
    match = DOCUMENT_ASSET_URL_RE.match(path)
    if match is None:
        return None
    try:
        document_id = uuid.UUID(match.group("document_id"))
        asset_id = uuid.UUID(match.group("asset_id"))
    except ValueError:
        return None
    if document_id != context.document_id:
        return None
    filename = _safe_asset_filename(unquote(match.group("filename")), f"{asset_id}")
    return document_asset_key(
        context.user_id,
        context.matter_id,
        context.document_id,
        asset_id,
        filename,
    )


def _add_tiptap_image(paragraph, attrs: dict[str, Any], context: DocumentAssetContext | None) -> bool:
    key = _document_asset_key_from_src(
        attrs.get("src") if isinstance(attrs.get("src"), str) else None,
        context,
    )
    if key is None:
        return False
    try:
        data = get_storage_backend().get_bytes(key)
    except (KeyError, StorageReadError):
        return False
    try:
        paragraph.add_run().add_picture(io.BytesIO(data), width=Inches(5.8))
    except Exception:
        return False
    return True


def _mime_from_filename(filename: str | None) -> str:
    suffix = filename.rsplit(".", 1)[-1].lower() if filename and "." in filename else ""
    return {
        "gif": "image/gif",
        "jpg": "image/jpeg",
        "jpeg": "image/jpeg",
        "png": "image/png",
        "webp": "image/webp",
    }.get(suffix, "application/octet-stream")


def _tiptap_image_data_uri(attrs: dict[str, Any], context: DocumentAssetContext | None) -> str | None:
    key = _document_asset_key_from_src(
        attrs.get("src") if isinstance(attrs.get("src"), str) else None,
        context,
    )
    if key is None:
        return None
    try:
        data = get_storage_backend().get_bytes(key)
    except (KeyError, StorageReadError):
        return None
    mime_type = _mime_from_filename(attrs.get("src") if isinstance(attrs.get("src"), str) else key)
    encoded = base64.b64encode(data).decode("ascii")
    return f"data:{mime_type};base64,{encoded}"


def _add_tiptap_inline_content(
    paragraph,
    nodes: list[dict[str, Any]] | None,
    asset_context: DocumentAssetContext | None = None,
) -> None:
    for node in nodes or []:
        node_type = node.get("type")
        if node_type == "text":
            run = paragraph.add_run(str(node.get("text") or ""))
            mark_types = {
                str(mark.get("type"))
                for mark in node.get("marks") or []
                if isinstance(mark, dict)
            }
            run.bold = "bold" in mark_types
            run.italic = "italic" in mark_types
            run.underline = "underline" in mark_types
            if "highlight" in mark_types:
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            for mark in node.get("marks") or []:
                if not isinstance(mark, dict) or mark.get("type") != "textStyle":
                    continue
                color = _rgb_from_hex((mark.get("attrs") or {}).get("color"))
                if color is not None:
                    run.font.color.rgb = color
        elif node_type == "hardBreak":
            paragraph.add_run().add_break()
        elif node_type == "image":
            attrs = node.get("attrs") or {}
            if not _add_tiptap_image(paragraph, attrs, asset_context):
                label = attrs.get("alt") or attrs.get("src") or "image"
                run = paragraph.add_run(f"[image: {label}]")
                run.italic = True
        elif isinstance(node.get("content"), list):
            _add_tiptap_inline_content(paragraph, node.get("content"), asset_context)


def _add_tiptap_paragraph(
    document: DocxDocument,
    node: dict[str, Any],
    style: str | None = None,
    asset_context: DocumentAssetContext | None = None,
) -> None:
    if node.get("type") == "heading":
        level = int((node.get("attrs") or {}).get("level") or 1)
        paragraph = document.add_heading(level=max(1, min(level, 4)))
    else:
        paragraph = document.add_paragraph(style=style)
    alignment = _paragraph_alignment((node.get("attrs") or {}).get("textAlign"))
    if alignment is not None:
        paragraph.alignment = alignment
    _add_tiptap_inline_content(paragraph, node.get("content"), asset_context)


def _add_tiptap_list_item(
    document: DocxDocument,
    item: dict[str, Any],
    style: str,
    asset_context: DocumentAssetContext | None = None,
) -> None:
    children = item.get("content") or []
    wrote_paragraph = False
    for child in children:
        if not isinstance(child, dict):
            continue
        child_type = child.get("type")
        if child_type in {"paragraph", "heading"}:
            _add_tiptap_paragraph(
                document,
                child,
                style=style if not wrote_paragraph else None,
                asset_context=asset_context,
            )
            wrote_paragraph = True
        elif child_type == "bulletList":
            _add_tiptap_list(document, child, "List Bullet", asset_context)
        elif child_type == "orderedList":
            _add_tiptap_list(document, child, "List Number", asset_context)
    if not wrote_paragraph:
        document.add_paragraph(style=style)


def _add_tiptap_list(
    document: DocxDocument,
    node: dict[str, Any],
    style: str,
    asset_context: DocumentAssetContext | None = None,
) -> None:
    for child in node.get("content") or []:
        if isinstance(child, dict) and child.get("type") == "listItem":
            _add_tiptap_list_item(document, child, style, asset_context)


def _add_tiptap_task_item(
    document: DocxDocument,
    item: dict[str, Any],
    asset_context: DocumentAssetContext | None = None,
) -> None:
    checked = bool((item.get("attrs") or {}).get("checked"))
    prefix = "[x] " if checked else "[ ] "
    children = item.get("content") or []
    paragraph_children: list[dict[str, Any]] | None = None
    for child in children:
        if isinstance(child, dict) and child.get("type") == "paragraph":
            paragraph_children = child.get("content")
            break
    paragraph = document.add_paragraph()
    paragraph.add_run(prefix)
    _add_tiptap_inline_content(paragraph, paragraph_children, asset_context)


def _add_tiptap_task_list(
    document: DocxDocument,
    node: dict[str, Any],
    asset_context: DocumentAssetContext | None = None,
) -> None:
    for child in node.get("content") or []:
        if isinstance(child, dict) and child.get("type") == "taskItem":
            _add_tiptap_task_item(document, child, asset_context)


def _tiptap_cell_text(cell: dict[str, Any]) -> str:
    text = _plain_text_from_tiptap_node(cell).strip()
    return re.sub(r"\n{2,}", "\n", text)


def _plain_text_from_tiptap_node(node: dict[str, Any]) -> str:
    node_type = node.get("type")
    if node_type == "text":
        return str(node.get("text") or "")
    if node_type == "hardBreak":
        return "\n"
    if node_type == "image":
        attrs = node.get("attrs") or {}
        label = attrs.get("alt") or attrs.get("src") or "image"
        return f"[image: {label}]\n"
    children = "".join(
        _plain_text_from_tiptap_node(child)
        for child in node.get("content") or []
        if isinstance(child, dict)
    )
    if node_type == "taskItem":
        prefix = "[x] " if (node.get("attrs") or {}).get("checked") else "[ ] "
        return f"{prefix}{children.strip()}\n"
    if node_type in {"paragraph", "heading", "listItem"}:
        return f"{children}\n"
    if node_type == "tableRow":
        return "\t".join(
            _plain_text_from_tiptap_node(child).strip()
            for child in node.get("content") or []
            if isinstance(child, dict)
        )
    return children


def _add_tiptap_table(document: DocxDocument, node: dict[str, Any]) -> None:
    rows = [
        row
        for row in node.get("content") or []
        if isinstance(row, dict) and row.get("type") == "tableRow"
    ]
    if not rows:
        return
    col_count = max(
        len(
            [
                cell
                for cell in row.get("content") or []
                if isinstance(cell, dict) and cell.get("type") in {"tableCell", "tableHeader"}
            ]
        )
        for row in rows
    )
    if col_count == 0:
        return
    table = document.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"
    for row_idx, row in enumerate(rows):
        cells = [
            cell
            for cell in row.get("content") or []
            if isinstance(cell, dict) and cell.get("type") in {"tableCell", "tableHeader"}
        ]
        for col_idx, cell in enumerate(cells[:col_count]):
            table_cell = table.rows[row_idx].cells[col_idx]
            table_cell.text = _tiptap_cell_text(cell)
            if cell.get("type") == "tableHeader":
                for para in table_cell.paragraphs:
                    for run in para.runs:
                        run.bold = True


def _html_attrs(attrs: dict[str, str]) -> str:
    pairs = [
        f'{name}="{html.escape(value, quote=True)}"'
        for name, value in attrs.items()
        if value
    ]
    return " " + " ".join(pairs) if pairs else ""


def _tiptap_inline_html(
    nodes: list[dict[str, Any]] | None,
    asset_context: DocumentAssetContext | None = None,
) -> str:
    parts: list[str] = []
    for node in nodes or []:
        node_type = node.get("type")
        if node_type == "text":
            text = html.escape(str(node.get("text") or ""))
            for mark in node.get("marks") or []:
                if not isinstance(mark, dict):
                    continue
                mark_type = mark.get("type")
                if mark_type == "bold":
                    text = f"<strong>{text}</strong>"
                elif mark_type == "italic":
                    text = f"<em>{text}</em>"
                elif mark_type == "underline":
                    text = f"<u>{text}</u>"
                elif mark_type == "highlight":
                    text = f"<mark>{text}</mark>"
                elif mark_type == "textStyle":
                    color = (mark.get("attrs") or {}).get("color")
                    if _rgb_from_hex(color) is not None:
                        text = f'<span style="color:{html.escape(str(color))}">{text}</span>'
            parts.append(text)
        elif node_type == "hardBreak":
            parts.append("<br/>")
        elif node_type == "image":
            attrs = node.get("attrs") or {}
            alt = html.escape(str(attrs.get("alt") or "image"))
            data_uri = _tiptap_image_data_uri(attrs, asset_context)
            if data_uri:
                parts.append(f'<img src="{data_uri}" alt="{alt}" />')
            else:
                parts.append(f'<em class="image-placeholder">[image: {alt}]</em>')
        elif isinstance(node.get("content"), list):
            parts.append(_tiptap_inline_html(node.get("content"), asset_context))
    return "".join(parts)


def _tiptap_block_html(
    node: dict[str, Any],
    asset_context: DocumentAssetContext | None = None,
) -> str:
    node_type = node.get("type")
    attrs = node.get("attrs") or {}
    align = attrs.get("textAlign") if attrs.get("textAlign") in {"left", "center", "right"} else None
    align_attr = _html_attrs({"style": f"text-align:{align}"} if align else {})
    if node_type == "paragraph":
        return f"<p{align_attr}>{_tiptap_inline_html(node.get('content'), asset_context)}</p>"
    if node_type == "heading":
        level = max(1, min(int(attrs.get("level") or 2), 4))
        return f"<h{level}{align_attr}>{_tiptap_inline_html(node.get('content'), asset_context)}</h{level}>"
    if node_type == "bulletList":
        items = "".join(
            f"<li>{_tiptap_inline_html(child.get('content'), asset_context)}</li>"
            if isinstance(child, dict) and child.get("type") == "listItem"
            else ""
            for child in node.get("content") or []
        )
        return f"<ul>{items}</ul>"
    if node_type == "orderedList":
        items = "".join(
            f"<li>{_tiptap_inline_html(child.get('content'), asset_context)}</li>"
            if isinstance(child, dict) and child.get("type") == "listItem"
            else ""
            for child in node.get("content") or []
        )
        return f"<ol>{items}</ol>"
    if node_type == "taskList":
        items = []
        for child in node.get("content") or []:
            if not isinstance(child, dict) or child.get("type") != "taskItem":
                continue
            checked = "x" if (child.get("attrs") or {}).get("checked") else " "
            items.append(
                f"<li><span class=\"task-box\">[{checked}]</span> "
                f"{_tiptap_inline_html(child.get('content'), asset_context)}</li>"
            )
        return f"<ul class=\"task-list\">{''.join(items)}</ul>"
    if node_type == "image":
        return f"<figure>{_tiptap_inline_html([node], asset_context)}</figure>"
    if node_type == "table":
        rows = []
        for row in node.get("content") or []:
            if not isinstance(row, dict) or row.get("type") != "tableRow":
                continue
            cells = []
            for cell in row.get("content") or []:
                if not isinstance(cell, dict):
                    continue
                tag = "th" if cell.get("type") == "tableHeader" else "td"
                cells.append(f"<{tag}>{html.escape(_tiptap_cell_text(cell))}</{tag}>")
            rows.append(f"<tr>{''.join(cells)}</tr>")
        return f"<table>{''.join(rows)}</table>"
    return ""


def _render_tiptap_html_body(
    body_json: dict[str, Any] | None,
    fallback_text: str,
    asset_context: DocumentAssetContext | None = None,
) -> str:
    if (
        not body_json
        or body_json.get("type") != "doc"
        or not isinstance(body_json.get("content"), list)
    ):
        return "".join(
            f"<p>{html.escape(part)}</p>"
            for part in re.split(r"\n{2,}", fallback_text)
            if part.strip()
        )
    return "".join(
        _tiptap_block_html(node, asset_context)
        for node in body_json.get("content") or []
        if isinstance(node, dict)
    )


def _render_document_version_html(
    title: str,
    version: DocumentVersion,
    comments: list[DocumentComment],
    asset_context: DocumentAssetContext,
) -> str:
    body_html = _render_tiptap_html_body(
        version.resolved_json,
        version.resolved_text or "",
        asset_context,
    )
    comments_html = "".join(
        f"<li><strong>{html.escape(comment.status)}</strong> "
        f"{html.escape(comment.body)}</li>"
        for comment in comments
    )
    comments_section = (
        f"<section><h2>Review notes</h2><ul>{comments_html}</ul></section>"
        if comments_html
        else ""
    )
    rendered_at = datetime.now(UTC).strftime("%Y-%m-%d %H:%M UTC")
    return f"""<!doctype html>
<html lang="en">
<head>
<meta charset="utf-8" />
<title>{html.escape(title)} v{version.version_number}</title>
<style>
  @page {{ size: A4; margin: 22mm 18mm; }}
  body {{ font-family: Inter, Helvetica, Arial, sans-serif; font-size: 11pt; line-height: 1.5; color: #1f2124; }}
  h1 {{ font-size: 22pt; margin: 0 0 8pt; }}
  h2 {{ font-size: 14pt; margin: 18pt 0 6pt; border-bottom: 1px solid #d4d4d4; padding-bottom: 4pt; }}
  h3 {{ font-size: 12pt; margin: 14pt 0 5pt; }}
  .meta {{ font-family: "Courier New", monospace; font-size: 9pt; color: #5d5e61; margin-bottom: 16pt; }}
  p {{ margin: 0 0 9pt; }}
  mark {{ background: #fff1a8; padding: 0 1pt; }}
  img {{ max-width: 100%; height: auto; margin: 8pt 0; }}
  figure {{ margin: 10pt 0; }}
  table {{ border-collapse: collapse; width: 100%; margin: 10pt 0; font-size: 10pt; }}
  th, td {{ border: 1px solid #d4d4d4; padding: 5pt 7pt; text-align: left; vertical-align: top; }}
  th {{ background: #f3f3f3; font-weight: 600; }}
  ul, ol {{ padding-left: 18pt; }}
  li {{ margin-bottom: 5pt; }}
  .task-list {{ list-style: none; padding-left: 0; }}
  .task-box {{ font-family: "Courier New", monospace; color: #5d5e61; }}
  .image-placeholder {{ color: #5d5e61; }}
  footer {{ margin-top: 24pt; padding-top: 6pt; border-top: 1px solid #d4d4d4; font-family: "Courier New", monospace; font-size: 9pt; color: #5d5e61; }}
</style>
</head>
<body>
  <h1>{html.escape(title)}</h1>
  <div class="meta">version: {version.version_number} · rendered: {rendered_at}</div>
  <main>{body_html}</main>
  {comments_section}
  <footer>Legalise document export · {rendered_at}</footer>
</body>
</html>"""


async def _html_to_pdf(html_doc: str) -> bytes:
    files = {"files": ("index.html", html_doc, "text/html")}
    url = f"{settings.gotenberg_url.rstrip('/')}/forms/chromium/convert/html"
    try:
        async with httpx.AsyncClient(timeout=60.0) as client:
            resp = await client.post(url, files=files)
    except httpx.HTTPError as exc:
        raise RuntimeError(f"Gotenberg unreachable at {url}: {exc}") from exc
    if resp.status_code != 200:
        raise RuntimeError(f"Gotenberg returned {resp.status_code}: {resp.text[:200]}")
    return resp.content


def _render_tiptap_docx(
    title: str,
    body_json: dict[str, Any],
    fallback_text: str,
    comments: list[DocumentComment] | None = None,
    asset_context: DocumentAssetContext | None = None,
) -> bytes:
    if body_json.get("type") != "doc" or not isinstance(body_json.get("content"), list):
        return _render_resolved_text_docx(title, fallback_text, comments or [])

    document = DocxDocument()
    document.add_heading(title, level=0)
    for node in body_json.get("content") or []:
        if not isinstance(node, dict):
            continue
        node_type = node.get("type")
        if node_type in {"paragraph", "heading"}:
            _add_tiptap_paragraph(document, node, asset_context=asset_context)
        elif node_type == "bulletList":
            _add_tiptap_list(document, node, "List Bullet", asset_context)
        elif node_type == "orderedList":
            _add_tiptap_list(document, node, "List Number", asset_context)
        elif node_type == "taskList":
            _add_tiptap_task_list(document, node, asset_context)
        elif node_type == "image":
            paragraph = document.add_paragraph()
            _add_tiptap_inline_content(paragraph, [node], asset_context)
        elif node_type == "table":
            _add_tiptap_table(document, node)
    _append_document_comments_docx(document, comments or [])
    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


# Route modules consume these via `from .common import *`; underscore-prefixed
# helpers must therefore be exported explicitly.
__all__ = [name for name in globals() if not name.startswith("__")]
