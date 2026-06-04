"""Documents API — per-document endpoints."""

from __future__ import annotations

import io
import base64
import hashlib
import html
import re
import uuid
from dataclasses import dataclass
from datetime import UTC, datetime, timedelta
from urllib.parse import unquote, urlparse

from typing import Any, Literal

import httpx
from fastapi import APIRouter, Depends, File, Form, HTTPException, Query, Response, UploadFile
from fastapi.responses import StreamingResponse
from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.shared import Inches, RGBColor
from pydantic import BaseModel, Field
from sqlalchemy import func, select
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.auth import current_user
from app.core.db import get_session
from app.core.document_uploads import (
    ALLOWED_UPLOAD_MIMES,
    MAX_UPLOAD_BYTES,
    MIME_TO_FORMAT,
    sniff_format,
)
from app.core.storage import (
    document_asset_key,
    get_storage_backend,
    StorageReadError,
    StorageWriteError,
    uploaded_key,
)
from app.core.text_extraction import extract as extract_text
from app.core.model_gateway import gateway as model_gateway
from app.core.api import PROVIDER_HTTP_EXCEPTIONS, audit, audit_failure, provider_error_http_exception
from app.core.config import settings
from app.models import (
    AuditEntry,
    COMMENT_STATUS_OPEN,
    COMMENT_STATUS_RESOLVED,
    Document,
    DocumentComment,
    DocumentEdit,
    DocumentEditSession,
    DocumentVersion,
    DocumentWorkingDraft,
    Matter,
    STATUS_ARCHIVED,
    User,
)
from app.models.document_body import DocumentBody, BODY_KIND_EXTRACTED, extracted_body_for
from app.models.document_edit import (
    EDIT_STATUS_ACCEPTED,
    EDIT_STATUS_PENDING,
    EDIT_STATUS_REJECTED,
)
from app.models.document_version import (
    VERSION_KIND_RESTORED,
    VERSION_KIND_UPLOAD,
    VERSION_KIND_USER_EDIT,
)
from app.modules.document_edit import EDIT_MODES, propose_edits
from app.modules.document_edit.resolver import (
    EditAlreadyResolved,
    resolve_bulk,
    resolve_edit,
)
from app.models.document_body import BODY_KIND_REDACTED
from app.modules.anonymisation.pipeline import anonymise_document
from app.modules.anonymisation.schemas import (
    AnonymisationResult,
    AnonymiseRequest,
    MappingRead,
    TokenMapping,
)

router = APIRouter()

IMAGE_ASSET_MIMES = {
    "image/gif",
    "image/jpeg",
    "image/png",
    "image/webp",
}
MAX_DOCUMENT_ASSET_BYTES = 5 * 1024 * 1024
DOCUMENT_ASSET_URL_RE = re.compile(
    r"^/api/documents/"
    r"(?P<document_id>[0-9a-fA-F-]{36})/assets/"
    r"(?P<asset_id>[0-9a-fA-F-]{36})/"
    r"(?P<filename>[^?#]+)$"
)


@dataclass(frozen=True)
class DocumentAssetContext:
    user_id: uuid.UUID
    matter_id: uuid.UUID
    document_id: uuid.UUID


class DocumentAssetUploadRead(BaseModel):
    id: uuid.UUID
    filename: str
    mime_type: str
    size_bytes: int
    sha256: str
    url: str


class DocumentBodyRead(BaseModel):
    document_id: uuid.UUID
    kind: str
    extracted_text: str
    extraction_method: str
    extracted_at: datetime
    char_count: int
    page_count: int | None
    error_reason: str | None = None

    model_config = {"from_attributes": True}


@router.get("/{document_id}/body", response_model=DocumentBodyRead)
async def get_document_body(
    document_id: uuid.UUID,
    plugin: str | None = None,
    skill: str | None = None,
    session: AsyncSession = Depends(get_session),
    user: User = Depends(current_user),
) -> DocumentBody:
    """Return the extracted body of a document.

    Authorisation: 404 if the document isn't found or its matter isn't
    owned by the current user.

    Body row semantics:
      - Row exists with `extraction_method='failed'` → 200 with empty
        text and `error_reason` populated. UI can surface the failure.
      - No row at all → 404 (extraction never ran).
    """
    row = await session.execute(
        select(Document, Matter)
        .join(Matter, Matter.id == Document.matter_id)
        .where(Document.id == document_id)
    )
    pair = row.first()
    if pair is None:
        raise HTTPException(404, "document not found")

    doc, matter = pair
    if matter.created_by_id != user.id or matter.status == STATUS_ARCHIVED:
        raise HTTPException(404, "document not found")

    # Module-attributed reads require `document.body.read` for the
    # `(plugin, skill)` triple. User-initiated UI reads (no plugin/skill
    # query params) keep the existing owner-only gate above.
    if plugin and skill:
        from app.core.capabilities import require_capability

        await require_capability(
            session,
            user_id=user.id,
            plugin=plugin,
            skill=skill,
            capability="document.body.read",
        )

    body = await session.scalar(
        select(DocumentBody).where(
            DocumentBody.document_id == doc.id,
            DocumentBody.kind == BODY_KIND_EXTRACTED,
        )
    )
    if body is None:
        raise HTTPException(404, "document body not available")
    return body


# -- Edit-instruction surface (Workstream 2) -------------------------------


EditMode = Literal["tighten", "rewrite", "summarise", "free-text", "uk-jurisdiction-sweep"]


class EditInstructionRequest(BaseModel):
    instruction: str = Field(min_length=4, max_length=2000)
    mode: EditMode = "free-text"


class DocumentVersionRead(BaseModel):
    id: uuid.UUID
    document_id: uuid.UUID
    version_number: int
    kind: str
    created_by_id: uuid.UUID
    created_at: datetime
    storage_uri: str | None
    filename: str | None = None
    mime_type: str | None = None
    size_bytes: int | None = None
    sha256: str | None = None
    notes: str | None
    resolved_text: str | None = None
    resolved_json: dict[str, Any] | None = None

    model_config = {"from_attributes": True}


class DocumentEditRead(BaseModel):
    id: uuid.UUID
    document_version_id: uuid.UUID
    change_id: str
    correlation_id: str | None
    deleted_text: str
    inserted_text: str
    context_before: str
    context_after: str
    rationale: str | None
    status: str
    created_at: datetime

    model_config = {"from_attributes": True}


class DocumentEditSessionStart(BaseModel):
    client_id: str = Field(min_length=8, max_length=96)


class DocumentEditSessionRead(BaseModel):
    id: uuid.UUID
    document_id: uuid.UUID
    user_id: uuid.UUID
    client_id: str
    user_label: str
    started_at: datetime
    last_seen_at: datetime
    ended_at: datetime | None


class DocumentEditSessionResponse(BaseModel):
    current: DocumentEditSessionRead
    active: list[DocumentEditSessionRead]


class DocumentWorkingDraftRead(BaseModel):
    document_id: uuid.UUID
    updated_by_id: uuid.UUID | None
    updated_at: datetime | None
    plain_text: str
    editor_json: dict[str, Any] | None = None
    base_version_id: uuid.UUID | None = None
    version_counter: int
    client_id: str | None = None

    model_config = {"from_attributes": True}


class DocumentWorkingDraftUpsert(BaseModel):
    plain_text: str = Field(max_length=500_000)
    editor_json: dict[str, Any] | None = None
    base_version_id: uuid.UUID | None = None
    client_id: str | None = Field(default=None, max_length=96)
    expected_version_counter: int | None = Field(default=None, ge=0)


class DocumentWorkingDraftCommitRequest(BaseModel):
    notes: str | None = Field(default=None, max_length=500)
    clear_draft: bool = True
    expected_version_counter: int | None = Field(default=None, ge=0)


class EditInstructionResponse(BaseModel):
    version: DocumentVersionRead
    pending_edits: list[DocumentEditRead]
    model_used: str
    model_notes: str
    instruction_hash: str
    parse_ok: bool


def _active_session_cutoff() -> datetime:
    return datetime.now(UTC) - timedelta(seconds=90)


async def _active_edit_sessions(
    session: AsyncSession,
    document_id: uuid.UUID,
) -> list[DocumentEditSessionRead]:
    rows = (
        await session.execute(
            select(DocumentEditSession, User)
            .join(User, User.id == DocumentEditSession.user_id)
            .where(
                DocumentEditSession.document_id == document_id,
                DocumentEditSession.ended_at.is_(None),
                DocumentEditSession.last_seen_at >= _active_session_cutoff(),
            )
            .order_by(DocumentEditSession.last_seen_at.desc())
        )
    ).all()
    return [
        DocumentEditSessionRead(
            id=row.id,
            document_id=row.document_id,
            user_id=row.user_id,
            client_id=row.client_id,
            user_label=user.name or user.email,
            started_at=row.started_at,
            last_seen_at=row.last_seen_at,
            ended_at=row.ended_at,
        )
        for row, user in rows
    ]


async def _latest_text_version(
    session: AsyncSession,
    document_id: uuid.UUID,
) -> DocumentVersion | None:
    return await session.scalar(
        select(DocumentVersion)
        .where(
            DocumentVersion.document_id == document_id,
            DocumentVersion.resolved_text.is_not(None),
        )
        .order_by(DocumentVersion.version_number.desc(), DocumentVersion.created_at.desc())
        .limit(1)
    )


async def _initial_working_draft(
    session: AsyncSession,
    doc: Document,
) -> DocumentWorkingDraftRead:
    latest = await _latest_text_version(session, doc.id)
    if latest is not None:
        return DocumentWorkingDraftRead(
            document_id=doc.id,
            updated_by_id=None,
            updated_at=None,
            plain_text=latest.resolved_text or "",
            editor_json=latest.resolved_json,
            base_version_id=latest.id,
            version_counter=0,
            client_id=None,
        )

    body = await extracted_body_for(session, doc.id)
    return DocumentWorkingDraftRead(
        document_id=doc.id,
        updated_by_id=None,
        updated_at=None,
        plain_text=body.extracted_text if body is not None else "",
        editor_json=None,
        base_version_id=None,
        version_counter=0,
        client_id=None,
    )


async def _create_user_edit_version(
    session: AsyncSession,
    *,
    doc: Document,
    matter: Matter,
    user: User,
    resolved_text: str,
    resolved_json: dict[str, Any] | None,
    notes: str | None,
    audit_source: str = "manual",
) -> DocumentVersion:
    next_version = (
        await session.scalar(
            select(func.coalesce(func.max(DocumentVersion.version_number), 0) + 1)
            .where(DocumentVersion.document_id == doc.id)
        )
        or 1
    )
    version = DocumentVersion(
        document_id=doc.id,
        version_number=int(next_version),
        kind=VERSION_KIND_USER_EDIT,
        created_by_id=user.id,
        filename=doc.filename,
        mime_type="text/plain",
        size_bytes=len(resolved_text.encode("utf-8")),
        sha256=hashlib.sha256(resolved_text.encode("utf-8")).hexdigest(),
        resolved_text=resolved_text,
        resolved_json=resolved_json,
        notes=notes or "Edited in Legalise document editor",
    )
    session.add(version)
    await session.flush()
    await audit.log(
        session,
        "document.version.saved",
        actor_id=user.id,
        matter_id=matter.id,
        module="document_editor",
        resource_type="document_version",
        resource_id=str(version.id),
        payload={
            "document_id": str(doc.id),
            "version_number": version.version_number,
            "kind": version.kind,
            "char_count": len(resolved_text),
            "rich_json": resolved_json is not None,
            "source": audit_source,
        },
    )
    return version


@router.post("/{document_id}/edit-instructions", response_model=EditInstructionResponse)
async def post_edit_instruction(
    document_id: uuid.UUID,
    body: EditInstructionRequest,
    session: AsyncSession = Depends(get_session),
    user: User = Depends(current_user),
) -> EditInstructionResponse:
    """Propose model edits to a document.

    Returns a new `assistant_edit` version + pending edits. The audit log carries
    a `module=document_edit, action=document.edit_instruction.invoked` row
    alongside the `model.call` row written by the gateway.
    """
    if body.mode not in EDIT_MODES:
        raise HTTPException(400, f"unknown mode: {body.mode}")

    try:
        result = await propose_edits(
            session=session,
            gateway=model_gateway,
            document_id=document_id,
            actor_id=user.id,
            instruction=body.instruction,
            mode=body.mode,
        )
    except LookupError as exc:
        raise HTTPException(404, str(exc)) from exc
    except ValueError as exc:
        raise HTTPException(422, str(exc)) from exc
    except PROVIDER_HTTP_EXCEPTIONS as exc:
        raise provider_error_http_exception(exc) from exc

    await session.commit()
    return EditInstructionResponse(
        version=DocumentVersionRead.model_validate(result.version),
        pending_edits=[DocumentEditRead.model_validate(e) for e in result.pending_edits],
        model_used=result.model_used,
        model_notes=result.model_notes,
        instruction_hash=result.instruction_hash,
        parse_ok=result.parse_ok,
    )


# -- Generated .docx download ----------------------------------------------


_FILENAME_SAFE_RE = re.compile(r"[^A-Za-z0-9._-]+")


def _safe_filename(title: str | None, file_uuid: str) -> str:
    """Derive a human-readable .docx filename from the audit `title` payload.

    Falls back to `generated-{uuid8}.docx` when title is empty or sanitises
    to nothing. The slugified title is bounded at 80 chars.
    """
    if title:
        cleaned = _FILENAME_SAFE_RE.sub("-", title).strip("-._")
        cleaned = cleaned[:80].rstrip("-._")
        if cleaned:
            return f"{cleaned}.docx"
    return f"generated-{file_uuid[:8]}.docx"


def _safe_asset_filename(filename: str | None, fallback: str) -> str:
    if filename:
        cleaned = _FILENAME_SAFE_RE.sub("-", filename).strip("-._")
        cleaned = cleaned[:100].rstrip("-._")
        if cleaned:
            return cleaned
    return fallback


def _docx_export_filename(filename: str, version_number: int) -> str:
    stem = filename.rsplit(".", 1)[0] if "." in filename else filename
    cleaned = _FILENAME_SAFE_RE.sub("-", stem).strip("-._")[:80].rstrip("-._")
    if not cleaned:
        cleaned = "document"
    return f"{cleaned}-v{version_number}.docx"


def _pdf_export_filename(filename: str, version_number: int) -> str:
    stem = filename.rsplit(".", 1)[0] if "." in filename else filename
    cleaned = _FILENAME_SAFE_RE.sub("-", stem).strip("-._")[:80].rstrip("-._")
    if not cleaned:
        cleaned = "document"
    return f"{cleaned}-v{version_number}.pdf"


async def _owned_live_document(
    session: AsyncSession,
    document_id: uuid.UUID,
    user: User,
) -> tuple[Document, Matter]:
    row = await session.execute(
        select(Document, Matter)
        .join(Matter, Matter.id == Document.matter_id)
        .where(Document.id == document_id)
    )
    pair = row.first()
    if pair is None:
        raise HTTPException(404, "document not found")
    doc, matter = pair
    if matter.created_by_id != user.id or matter.status == STATUS_ARCHIVED:
        raise HTTPException(404, "document not found")
    return doc, matter


@router.post("/{document_id}/assets", response_model=DocumentAssetUploadRead)
async def post_document_asset(
    document_id: uuid.UUID,
    file: UploadFile = File(...),
    session: AsyncSession = Depends(get_session),
    user: User = Depends(current_user),
) -> DocumentAssetUploadRead:
    """Upload an embedded editor image for a document.

    Assets live under the matter storage prefix and are retrieved through
    the backend, so auth and matter cleanup stay inside the existing
    document boundary. Reads are not audited because every render would
    otherwise create noisy rows; the upload itself is recorded.
    """
    doc, matter = await _owned_live_document(session, document_id, user)
    mime_type = file.content_type or "application/octet-stream"
    if mime_type not in IMAGE_ASSET_MIMES:
        raise HTTPException(415, "unsupported image type")
    data = await file.read()
    if not data:
        raise HTTPException(400, "empty image")
    if len(data) > MAX_DOCUMENT_ASSET_BYTES:
        raise HTTPException(413, "image too large")

    asset_id = uuid.uuid4()
    extension = {
        "image/gif": "gif",
        "image/jpeg": "jpg",
        "image/png": "png",
        "image/webp": "webp",
    }[mime_type]
    filename = _safe_asset_filename(file.filename, f"{asset_id}.{extension}")
    if "." not in filename:
        filename = f"{filename}.{extension}"
    digest = hashlib.sha256(data).hexdigest()
    key = document_asset_key(user.id, matter.id, doc.id, asset_id, filename)
    try:
        get_storage_backend().put_bytes(
            key,
            data,
            content_type=mime_type,
            metadata={"sha256": digest, "document_id": str(doc.id)},
        )
    except StorageWriteError as exc:
        await audit_failure(
            session,
            "storage.put_bytes.failed",
            actor_id=user.id,
            matter_id=matter.id,
            module="storage",
            resource_type="document_asset",
            resource_id=str(asset_id),
            payload={
                "storage_key": key,
                "backend": exc.backend,
                "error_code": exc.error_code,
            },
        )
        raise HTTPException(
            502,
            detail={
                "error": "storage_write_failed",
                "message": "Failed to store document image.",
                "storage_key": key,
                "backend": exc.backend,
            },
        ) from exc

    await audit.log(
        session,
        "document.asset.uploaded",
        actor_id=user.id,
        matter_id=matter.id,
        module="document_editor",
        resource_type="document_asset",
        resource_id=str(asset_id),
        payload={
            "document_id": str(doc.id),
            "filename": filename,
            "mime_type": mime_type,
            "size_bytes": len(data),
            "sha256": digest,
        },
    )
    await session.commit()
    return DocumentAssetUploadRead(
        id=asset_id,
        filename=filename,
        mime_type=mime_type,
        size_bytes=len(data),
        sha256=digest,
        url=f"/api/documents/{doc.id}/assets/{asset_id}/{filename}",
    )


@router.get("/{document_id}/assets/{asset_id}/{filename}")
async def get_document_asset(
    document_id: uuid.UUID,
    asset_id: uuid.UUID,
    filename: str,
    session: AsyncSession = Depends(get_session),
    user: User = Depends(current_user),
) -> StreamingResponse:
    doc, matter = await _owned_live_document(session, document_id, user)
    safe_filename = _safe_asset_filename(filename, f"{asset_id}")
    key = document_asset_key(user.id, matter.id, doc.id, asset_id, safe_filename)
    try:
        data = get_storage_backend().get_bytes(key)
    except KeyError:
        raise HTTPException(404, "document asset not found")
    except StorageReadError as exc:
        await audit_failure(
            session,
            "storage.get_bytes.failed",
            actor_id=user.id,
            matter_id=matter.id,
            module="storage",
            resource_type="document_asset",
            resource_id=str(asset_id),
            payload={
                "storage_key": key,
                "backend": exc.backend,
                "error_code": exc.error_code,
            },
        )
        raise HTTPException(
            502,
            detail={
                "error": "storage_read_failed",
                "message": "Failed to read document image.",
                "storage_key": key,
                "backend": exc.backend,
            },
        ) from exc

    suffix = safe_filename.rsplit(".", 1)[-1].lower() if "." in safe_filename else ""
    mime_type = {
        "gif": "image/gif",
        "jpg": "image/jpeg",
        "jpeg": "image/jpeg",
        "png": "image/png",
        "webp": "image/webp",
    }.get(suffix, "application/octet-stream")
    return StreamingResponse(
        iter([data]),
        media_type=mime_type,
        headers={"Content-Length": str(len(data))},
    )


def _render_resolved_text_docx(
    title: str,
    body: str,
    comments: list[DocumentComment] | None = None,
) -> bytes:
    document = DocxDocument()
    document.add_heading(title, level=0)
    for block in body.split("\n\n"):
        block = block.rstrip()
        if not block:
            continue
        lines = block.split("\n")
        para = document.add_paragraph(lines[0])
        for line in lines[1:]:
            para.add_run().add_break()
            para.add_run(line)
    buf = io.BytesIO()
    _append_document_comments_docx(document, comments or [])
    document.save(buf)
    return buf.getvalue()


def _append_document_comments_docx(
    document: DocxDocument,
    comments: list[DocumentComment],
) -> None:
    if not comments:
        return
    document.add_page_break()
    document.add_heading("Document review notes", level=1)
    for index, comment in enumerate(comments, start=1):
        status = "resolved" if comment.status == COMMENT_STATUS_RESOLVED else "open"
        document.add_heading(f"Note {index} ({status})", level=2)
        if comment.quote_text:
            quote = document.add_paragraph(style="Intense Quote")
            quote.add_run(comment.quote_text)
        document.add_paragraph(comment.body)
        meta = document.add_paragraph()
        meta.add_run("Created: ").bold = True
        meta.add_run(comment.created_at.isoformat())
        if comment.resolved_at:
            meta.add_run(" · Resolved: ").bold = True
            meta.add_run(comment.resolved_at.isoformat())


def _rgb_from_hex(value: str | None) -> RGBColor | None:
    if not value:
        return None
    cleaned = value.strip().lstrip("#")
    if len(cleaned) != 6:
        return None
    try:
        return RGBColor(
            int(cleaned[0:2], 16),
            int(cleaned[2:4], 16),
            int(cleaned[4:6], 16),
        )
    except ValueError:
        return None


def _paragraph_alignment(value: str | None):
    return {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }.get(value or "")


def _document_asset_key_from_src(
    src: str | None,
    context: DocumentAssetContext | None,
) -> str | None:
    if not src or context is None:
        return None
    parsed = urlparse(src)
    path = parsed.path if parsed.scheme or parsed.netloc else src
    match = DOCUMENT_ASSET_URL_RE.match(path)
    if match is None:
        return None
    try:
        document_id = uuid.UUID(match.group("document_id"))
        asset_id = uuid.UUID(match.group("asset_id"))
    except ValueError:
        return None
    if document_id != context.document_id:
        return None
    filename = _safe_asset_filename(unquote(match.group("filename")), f"{asset_id}")
    return document_asset_key(
        context.user_id,
        context.matter_id,
        context.document_id,
        asset_id,
        filename,
    )


def _add_tiptap_image(paragraph, attrs: dict[str, Any], context: DocumentAssetContext | None) -> bool:
    key = _document_asset_key_from_src(
        attrs.get("src") if isinstance(attrs.get("src"), str) else None,
        context,
    )
    if key is None:
        return False
    try:
        data = get_storage_backend().get_bytes(key)
    except (KeyError, StorageReadError):
        return False
    try:
        paragraph.add_run().add_picture(io.BytesIO(data), width=Inches(5.8))
    except Exception:
        return False
    return True


def _mime_from_filename(filename: str | None) -> str:
    suffix = filename.rsplit(".", 1)[-1].lower() if filename and "." in filename else ""
    return {
        "gif": "image/gif",
        "jpg": "image/jpeg",
        "jpeg": "image/jpeg",
        "png": "image/png",
        "webp": "image/webp",
    }.get(suffix, "application/octet-stream")


def _tiptap_image_data_uri(attrs: dict[str, Any], context: DocumentAssetContext | None) -> str | None:
    key = _document_asset_key_from_src(
        attrs.get("src") if isinstance(attrs.get("src"), str) else None,
        context,
    )
    if key is None:
        return None
    try:
        data = get_storage_backend().get_bytes(key)
    except (KeyError, StorageReadError):
        return None
    mime_type = _mime_from_filename(attrs.get("src") if isinstance(attrs.get("src"), str) else key)
    encoded = base64.b64encode(data).decode("ascii")
    return f"data:{mime_type};base64,{encoded}"


def _add_tiptap_inline_content(
    paragraph,
    nodes: list[dict[str, Any]] | None,
    asset_context: DocumentAssetContext | None = None,
) -> None:
    for node in nodes or []:
        node_type = node.get("type")
        if node_type == "text":
            run = paragraph.add_run(str(node.get("text") or ""))
            mark_types = {
                str(mark.get("type"))
                for mark in node.get("marks") or []
                if isinstance(mark, dict)
            }
            run.bold = "bold" in mark_types
            run.italic = "italic" in mark_types
            run.underline = "underline" in mark_types
            if "highlight" in mark_types:
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            for mark in node.get("marks") or []:
                if not isinstance(mark, dict) or mark.get("type") != "textStyle":
                    continue
                color = _rgb_from_hex((mark.get("attrs") or {}).get("color"))
                if color is not None:
                    run.font.color.rgb = color
        elif node_type == "hardBreak":
            paragraph.add_run().add_break()
        elif node_type == "image":
            attrs = node.get("attrs") or {}
            if not _add_tiptap_image(paragraph, attrs, asset_context):
                label = attrs.get("alt") or attrs.get("src") or "image"
                run = paragraph.add_run(f"[image: {label}]")
                run.italic = True
        elif isinstance(node.get("content"), list):
            _add_tiptap_inline_content(paragraph, node.get("content"), asset_context)


def _add_tiptap_paragraph(
    document: DocxDocument,
    node: dict[str, Any],
    style: str | None = None,
    asset_context: DocumentAssetContext | None = None,
) -> None:
    if node.get("type") == "heading":
        level = int((node.get("attrs") or {}).get("level") or 1)
        paragraph = document.add_heading(level=max(1, min(level, 4)))
    else:
        paragraph = document.add_paragraph(style=style)
    alignment = _paragraph_alignment((node.get("attrs") or {}).get("textAlign"))
    if alignment is not None:
        paragraph.alignment = alignment
    _add_tiptap_inline_content(paragraph, node.get("content"), asset_context)


def _add_tiptap_list_item(
    document: DocxDocument,
    item: dict[str, Any],
    style: str,
    asset_context: DocumentAssetContext | None = None,
) -> None:
    children = item.get("content") or []
    wrote_paragraph = False
    for child in children:
        if not isinstance(child, dict):
            continue
        child_type = child.get("type")
        if child_type in {"paragraph", "heading"}:
            _add_tiptap_paragraph(
                document,
                child,
                style=style if not wrote_paragraph else None,
                asset_context=asset_context,
            )
            wrote_paragraph = True
        elif child_type == "bulletList":
            _add_tiptap_list(document, child, "List Bullet", asset_context)
        elif child_type == "orderedList":
            _add_tiptap_list(document, child, "List Number", asset_context)
    if not wrote_paragraph:
        document.add_paragraph(style=style)


def _add_tiptap_list(
    document: DocxDocument,
    node: dict[str, Any],
    style: str,
    asset_context: DocumentAssetContext | None = None,
) -> None:
    for child in node.get("content") or []:
        if isinstance(child, dict) and child.get("type") == "listItem":
            _add_tiptap_list_item(document, child, style, asset_context)


def _add_tiptap_task_item(
    document: DocxDocument,
    item: dict[str, Any],
    asset_context: DocumentAssetContext | None = None,
) -> None:
    checked = bool((item.get("attrs") or {}).get("checked"))
    prefix = "[x] " if checked else "[ ] "
    children = item.get("content") or []
    paragraph_children: list[dict[str, Any]] | None = None
    for child in children:
        if isinstance(child, dict) and child.get("type") == "paragraph":
            paragraph_children = child.get("content")
            break
    paragraph = document.add_paragraph()
    paragraph.add_run(prefix)
    _add_tiptap_inline_content(paragraph, paragraph_children, asset_context)


def _add_tiptap_task_list(
    document: DocxDocument,
    node: dict[str, Any],
    asset_context: DocumentAssetContext | None = None,
) -> None:
    for child in node.get("content") or []:
        if isinstance(child, dict) and child.get("type") == "taskItem":
            _add_tiptap_task_item(document, child, asset_context)


def _tiptap_cell_text(cell: dict[str, Any]) -> str:
    text = _plain_text_from_tiptap_node(cell).strip()
    return re.sub(r"\n{2,}", "\n", text)


def _plain_text_from_tiptap_node(node: dict[str, Any]) -> str:
    node_type = node.get("type")
    if node_type == "text":
        return str(node.get("text") or "")
    if node_type == "hardBreak":
        return "\n"
    if node_type == "image":
        attrs = node.get("attrs") or {}
        label = attrs.get("alt") or attrs.get("src") or "image"
        return f"[image: {label}]\n"
    children = "".join(
        _plain_text_from_tiptap_node(child)
        for child in node.get("content") or []
        if isinstance(child, dict)
    )
    if node_type == "taskItem":
        prefix = "[x] " if (node.get("attrs") or {}).get("checked") else "[ ] "
        return f"{prefix}{children.strip()}\n"
    if node_type in {"paragraph", "heading", "listItem"}:
        return f"{children}\n"
    if node_type == "tableRow":
        return "\t".join(
            _plain_text_from_tiptap_node(child).strip()
            for child in node.get("content") or []
            if isinstance(child, dict)
        )
    return children


def _add_tiptap_table(document: DocxDocument, node: dict[str, Any]) -> None:
    rows = [
        row
        for row in node.get("content") or []
        if isinstance(row, dict) and row.get("type") == "tableRow"
    ]
    if not rows:
        return
    col_count = max(
        len(
            [
                cell
                for cell in row.get("content") or []
                if isinstance(cell, dict) and cell.get("type") in {"tableCell", "tableHeader"}
            ]
        )
        for row in rows
    )
    if col_count == 0:
        return
    table = document.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"
    for row_idx, row in enumerate(rows):
        cells = [
            cell
            for cell in row.get("content") or []
            if isinstance(cell, dict) and cell.get("type") in {"tableCell", "tableHeader"}
        ]
        for col_idx, cell in enumerate(cells[:col_count]):
            table_cell = table.rows[row_idx].cells[col_idx]
            table_cell.text = _tiptap_cell_text(cell)
            if cell.get("type") == "tableHeader":
                for para in table_cell.paragraphs:
                    for run in para.runs:
                        run.bold = True


def _html_attrs(attrs: dict[str, str]) -> str:
    pairs = [
        f'{name}="{html.escape(value, quote=True)}"'
        for name, value in attrs.items()
        if value
    ]
    return " " + " ".join(pairs) if pairs else ""


def _tiptap_inline_html(
    nodes: list[dict[str, Any]] | None,
    asset_context: DocumentAssetContext | None = None,
) -> str:
    parts: list[str] = []
    for node in nodes or []:
        node_type = node.get("type")
        if node_type == "text":
            text = html.escape(str(node.get("text") or ""))
            for mark in node.get("marks") or []:
                if not isinstance(mark, dict):
                    continue
                mark_type = mark.get("type")
                if mark_type == "bold":
                    text = f"<strong>{text}</strong>"
                elif mark_type == "italic":
                    text = f"<em>{text}</em>"
                elif mark_type == "underline":
                    text = f"<u>{text}</u>"
                elif mark_type == "highlight":
                    text = f"<mark>{text}</mark>"
                elif mark_type == "textStyle":
                    color = (mark.get("attrs") or {}).get("color")
                    if _rgb_from_hex(color) is not None:
                        text = f'<span style="color:{html.escape(str(color))}">{text}</span>'
            parts.append(text)
        elif node_type == "hardBreak":
            parts.append("<br/>")
        elif node_type == "image":
            attrs = node.get("attrs") or {}
            alt = html.escape(str(attrs.get("alt") or "image"))
            data_uri = _tiptap_image_data_uri(attrs, asset_context)
            if data_uri:
                parts.append(f'<img src="{data_uri}" alt="{alt}" />')
            else:
                parts.append(f'<em class="image-placeholder">[image: {alt}]</em>')
        elif isinstance(node.get("content"), list):
            parts.append(_tiptap_inline_html(node.get("content"), asset_context))
    return "".join(parts)


def _tiptap_block_html(
    node: dict[str, Any],
    asset_context: DocumentAssetContext | None = None,
) -> str:
    node_type = node.get("type")
    attrs = node.get("attrs") or {}
    align = attrs.get("textAlign") if attrs.get("textAlign") in {"left", "center", "right"} else None
    align_attr = _html_attrs({"style": f"text-align:{align}"} if align else {})
    if node_type == "paragraph":
        return f"<p{align_attr}>{_tiptap_inline_html(node.get('content'), asset_context)}</p>"
    if node_type == "heading":
        level = max(1, min(int(attrs.get("level") or 2), 4))
        return f"<h{level}{align_attr}>{_tiptap_inline_html(node.get('content'), asset_context)}</h{level}>"
    if node_type == "bulletList":
        items = "".join(
            f"<li>{_tiptap_inline_html(child.get('content'), asset_context)}</li>"
            if isinstance(child, dict) and child.get("type") == "listItem"
            else ""
            for child in node.get("content") or []
        )
        return f"<ul>{items}</ul>"
    if node_type == "orderedList":
        items = "".join(
            f"<li>{_tiptap_inline_html(child.get('content'), asset_context)}</li>"
            if isinstance(child, dict) and child.get("type") == "listItem"
            else ""
            for child in node.get("content") or []
        )
        return f"<ol>{items}</ol>"
    if node_type == "taskList":
        items = []
        for child in node.get("content") or []:
            if not isinstance(child, dict) or child.get("type") != "taskItem":
                continue
            checked = "x" if (child.get("attrs") or {}).get("checked") else " "
            items.append(
                f"<li><span class=\"task-box\">[{checked}]</span> "
                f"{_tiptap_inline_html(child.get('content'), asset_context)}</li>"
            )
        return f"<ul class=\"task-list\">{''.join(items)}</ul>"
    if node_type == "image":
        return f"<figure>{_tiptap_inline_html([node], asset_context)}</figure>"
    if node_type == "table":
        rows = []
        for row in node.get("content") or []:
            if not isinstance(row, dict) or row.get("type") != "tableRow":
                continue
            cells = []
            for cell in row.get("content") or []:
                if not isinstance(cell, dict):
                    continue
                tag = "th" if cell.get("type") == "tableHeader" else "td"
                cells.append(f"<{tag}>{html.escape(_tiptap_cell_text(cell))}</{tag}>")
            rows.append(f"<tr>{''.join(cells)}</tr>")
        return f"<table>{''.join(rows)}</table>"
    return ""


def _render_tiptap_html_body(
    body_json: dict[str, Any] | None,
    fallback_text: str,
    asset_context: DocumentAssetContext | None = None,
) -> str:
    if (
        not body_json
        or body_json.get("type") != "doc"
        or not isinstance(body_json.get("content"), list)
    ):
        return "".join(
            f"<p>{html.escape(part)}</p>"
            for part in re.split(r"\n{2,}", fallback_text)
            if part.strip()
        )
    return "".join(
        _tiptap_block_html(node, asset_context)
        for node in body_json.get("content") or []
        if isinstance(node, dict)
    )


def _render_document_version_html(
    title: str,
    version: DocumentVersion,
    comments: list[DocumentComment],
    asset_context: DocumentAssetContext,
) -> str:
    body_html = _render_tiptap_html_body(
        version.resolved_json,
        version.resolved_text or "",
        asset_context,
    )
    comments_html = "".join(
        f"<li><strong>{html.escape(comment.status)}</strong> "
        f"{html.escape(comment.body)}</li>"
        for comment in comments
    )
    comments_section = (
        f"<section><h2>Review notes</h2><ul>{comments_html}</ul></section>"
        if comments_html
        else ""
    )
    rendered_at = datetime.now(UTC).strftime("%Y-%m-%d %H:%M UTC")
    return f"""<!doctype html>
<html lang="en">
<head>
<meta charset="utf-8" />
<title>{html.escape(title)} v{version.version_number}</title>
<style>
  @page {{ size: A4; margin: 22mm 18mm; }}
  body {{ font-family: Inter, Helvetica, Arial, sans-serif; font-size: 11pt; line-height: 1.5; color: #1f2124; }}
  h1 {{ font-size: 22pt; margin: 0 0 8pt; }}
  h2 {{ font-size: 14pt; margin: 18pt 0 6pt; border-bottom: 1px solid #d4d4d4; padding-bottom: 4pt; }}
  h3 {{ font-size: 12pt; margin: 14pt 0 5pt; }}
  .meta {{ font-family: "Courier New", monospace; font-size: 9pt; color: #5d5e61; margin-bottom: 16pt; }}
  p {{ margin: 0 0 9pt; }}
  mark {{ background: #fff1a8; padding: 0 1pt; }}
  img {{ max-width: 100%; height: auto; margin: 8pt 0; }}
  figure {{ margin: 10pt 0; }}
  table {{ border-collapse: collapse; width: 100%; margin: 10pt 0; font-size: 10pt; }}
  th, td {{ border: 1px solid #d4d4d4; padding: 5pt 7pt; text-align: left; vertical-align: top; }}
  th {{ background: #f3f3f3; font-weight: 600; }}
  ul, ol {{ padding-left: 18pt; }}
  li {{ margin-bottom: 5pt; }}
  .task-list {{ list-style: none; padding-left: 0; }}
  .task-box {{ font-family: "Courier New", monospace; color: #5d5e61; }}
  .image-placeholder {{ color: #5d5e61; }}
  footer {{ margin-top: 24pt; padding-top: 6pt; border-top: 1px solid #d4d4d4; font-family: "Courier New", monospace; font-size: 9pt; color: #5d5e61; }}
</style>
</head>
<body>
  <h1>{html.escape(title)}</h1>
  <div class="meta">version: {version.version_number} · rendered: {rendered_at}</div>
  <main>{body_html}</main>
  {comments_section}
  <footer>Legalise document export · {rendered_at}</footer>
</body>
</html>"""


async def _html_to_pdf(html_doc: str) -> bytes:
    files = {"files": ("index.html", html_doc, "text/html")}
    url = f"{settings.gotenberg_url.rstrip('/')}/forms/chromium/convert/html"
    try:
        async with httpx.AsyncClient(timeout=60.0) as client:
            resp = await client.post(url, files=files)
    except httpx.HTTPError as exc:
        raise RuntimeError(f"Gotenberg unreachable at {url}: {exc}") from exc
    if resp.status_code != 200:
        raise RuntimeError(f"Gotenberg returned {resp.status_code}: {resp.text[:200]}")
    return resp.content


def _render_tiptap_docx(
    title: str,
    body_json: dict[str, Any],
    fallback_text: str,
    comments: list[DocumentComment] | None = None,
    asset_context: DocumentAssetContext | None = None,
) -> bytes:
    if body_json.get("type") != "doc" or not isinstance(body_json.get("content"), list):
        return _render_resolved_text_docx(title, fallback_text, comments or [])

    document = DocxDocument()
    document.add_heading(title, level=0)
    for node in body_json.get("content") or []:
        if not isinstance(node, dict):
            continue
        node_type = node.get("type")
        if node_type in {"paragraph", "heading"}:
            _add_tiptap_paragraph(document, node, asset_context=asset_context)
        elif node_type == "bulletList":
            _add_tiptap_list(document, node, "List Bullet", asset_context)
        elif node_type == "orderedList":
            _add_tiptap_list(document, node, "List Number", asset_context)
        elif node_type == "taskList":
            _add_tiptap_task_list(document, node, asset_context)
        elif node_type == "image":
            paragraph = document.add_paragraph()
            _add_tiptap_inline_content(paragraph, [node], asset_context)
        elif node_type == "table":
            _add_tiptap_table(document, node)
    _append_document_comments_docx(document, comments or [])
    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


@router.get("/generated/{file_uuid}")
async def download_generated_docx(
    file_uuid: uuid.UUID,
    session: AsyncSession = Depends(get_session),
    user: User = Depends(current_user),
) -> FileResponse:
    """Stream a previously generated .docx.

    Authorisation walks the audit trail: the canonical handle on a
    generated file is the `document.generated` AuditEntry written by the
    `generate_docx` tool. We resolve the most recent matching row, walk
    to its `matter_id`, and 404 unless `matter.created_by_id == user.id`.
    No row → 404. File missing on disk → 404 (treat as gone).
    """
    entry = await session.scalar(
        select(AuditEntry)
        .where(
            AuditEntry.action == "document.generated",
            AuditEntry.resource_id == str(file_uuid),
        )
        .order_by(AuditEntry.timestamp.desc())
        .limit(1)
    )
    if entry is None or entry.matter_id is None:
        raise HTTPException(404, "generated document not found")

    matter = await session.scalar(
        select(Matter).where(Matter.id == entry.matter_id)
    )
    if matter is None or matter.created_by_id != user.id or matter.status == STATUS_ARCHIVED:
        raise HTTPException(404, "generated document not found")

    storage_uri = (entry.payload or {}).get("storage_uri")
    if not storage_uri:
        raise HTTPException(404, "generated document not found")

    storage = get_storage_backend()
    try:
        data = storage.get_bytes(storage_uri)
    except KeyError:
        raise HTTPException(404, "generated document not found")
    except StorageReadError as exc:
        # Forensic provenance via `audit_failure` (separate committed
        # session) so the row survives the route's session rollback —
        # R3 review fix. Aligned with the upload-fail path.
        from app.core.api import audit_failure
        await audit_failure(
            session,
            "storage.get_bytes.failed",
            actor_id=user.id,
            matter_id=entry.matter_id,
            module="storage",
            resource_type="document",
            resource_id=str(file_uuid),
            payload={
                "storage_key": storage_uri,
                "backend": exc.backend,
                "error_code": exc.error_code,
            },
        )
        raise HTTPException(
            502,
            detail={
                "error": "storage_read_failed",
                "message": "Failed to read generated document from object storage.",
                "storage_key": storage_uri,
                "backend": exc.backend,
            },
        ) from exc

    filename = _safe_filename((entry.payload or {}).get("title"), str(file_uuid))
    return StreamingResponse(
        iter([data]),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f'attachment; filename="{filename}"',
            "Content-Length": str(len(data)),
        },
    )


# -- Original file retrieval (streamed backend proxy) ----------------------


@router.get("/{document_id}/original")
async def get_document_original(
    document_id: uuid.UUID,
    download: int = Query(0),
    session: AsyncSession = Depends(get_session),
    user: User = Depends(current_user),
) -> StreamingResponse:
    """Stream the original uploaded bytes for a matter document.

    Governed proxy (not a presigned URL): access stays behind the
    backend so auth, audit, and failure envelopes live inside the
    product boundary. **Owner-only** — matching the existing body /
    versions endpoints; there is deliberately no superuser/admin
    document-read shortcut (admin document inspection, if ever needed,
    must be a separate explicit policy, not smuggled into this path).
    Cross-user / archived / missing all return a uniform 404.
    `?download=1` switches the disposition from inline to attachment.
    Every successful access writes `document.original.accessed`.
    """
    doc = await session.scalar(select(Document).where(Document.id == document_id))
    if doc is None:
        raise HTTPException(404, "document not found")
    matter = await session.scalar(select(Matter).where(Matter.id == doc.matter_id))
    # Owner-only; archived is gone for everyone. Uniform 404 so we never
    # leak which documents exist for other users. No superuser branch by
    # design (see docstring).
    if (
        matter is None
        or matter.status == STATUS_ARCHIVED
        or matter.created_by_id != user.id
    ):
        raise HTTPException(404, "document not found")
    if not doc.storage_uri:
        raise HTTPException(404, "original file not available")

    storage = get_storage_backend()
    try:
        data = storage.get_bytes(doc.storage_uri)
    except KeyError:
        raise HTTPException(404, "original file not available")
    except StorageReadError as exc:
        await audit_failure(
            session,
            "storage.get_bytes.failed",
            actor_id=user.id,
            matter_id=doc.matter_id,
            module="storage",
            resource_type="document",
            resource_id=str(doc.id),
            payload={
                "storage_key": doc.storage_uri,
                "backend": exc.backend,
                "error_code": exc.error_code,
            },
        )
        raise HTTPException(
            502,
            detail={
                "error": "storage_read_failed",
                "message": "Failed to read the original document from object storage.",
                "storage_key": doc.storage_uri,
                "backend": exc.backend,
            },
        ) from exc

    is_download = bool(download)
    await audit.log(
        session,
        "document.original.accessed",
        actor_id=user.id,
        matter_id=doc.matter_id,
        resource_type="document",
        resource_id=str(doc.id),
        payload={
            "filename": doc.filename,
            "sha256": doc.sha256,
            "mime_type": doc.mime_type,
            "size_bytes": doc.size_bytes,
            "download": is_download,
        },
    )
    await session.commit()

    filename = _safe_filename(doc.filename, str(doc.id))
    disposition = "attachment" if is_download else "inline"
    return StreamingResponse(
        iter([data]),
        media_type=doc.mime_type or "application/octet-stream",
        headers={
            "Content-Disposition": f'{disposition}; filename="{filename}"',
            "Content-Length": str(len(data)),
        },
    )


# -- Accept/reject + versions (Workstream 2) -------------------------------


class EditResolutionResponse(BaseModel):
    edit: DocumentEditRead
    new_version: DocumentVersionRead | None = None
    resolved_text: str | None = None


class BulkResolutionResponse(BaseModel):
    affected_count: int
    new_version: DocumentVersionRead
    resolved_text: str


class DocumentVersionSummary(BaseModel):
    version: DocumentVersionRead
    pending_count: int
    accepted_count: int
    rejected_count: int


class ManualDocumentVersionRequest(BaseModel):
    resolved_text: str = Field(min_length=1, max_length=500_000)
    resolved_json: dict[str, Any] | None = None
    notes: str | None = Field(default=None, max_length=500)


class RestoreDocumentVersionRequest(BaseModel):
    notes: str | None = Field(default=None, max_length=500)


class DocumentCommentRead(BaseModel):
    id: uuid.UUID
    document_id: uuid.UUID
    author_id: uuid.UUID
    quote_text: str | None
    body_sha256: str | None
    anchor_start: int | None
    anchor_end: int | None
    body: str
    status: str
    created_at: datetime
    resolved_at: datetime | None
    resolved_by_id: uuid.UUID | None

    model_config = {"from_attributes": True}


class DocumentCommentCreate(BaseModel):
    body: str = Field(min_length=2, max_length=4000)
    quote_text: str | None = Field(default=None, max_length=2000)
    body_sha256: str | None = Field(default=None, max_length=64)
    anchor_start: int | None = Field(default=None, ge=0)
    anchor_end: int | None = Field(default=None, ge=0)


class DocumentCommentUpdate(BaseModel):
    body: str = Field(min_length=2, max_length=4000)


async def _resolve_one(
    document_id_unused: None,
    edit_id: uuid.UUID,
    action: Literal["accept", "reject"],
    session: AsyncSession,
    user: User,
) -> EditResolutionResponse:
    try:
        updated, new_version, resolved_text = await resolve_edit(
            session,
            edit_id=edit_id,
            actor_id=user.id,
            action=action,
        )
    except LookupError as exc:
        raise HTTPException(404, str(exc)) from exc
    except EditAlreadyResolved as exc:
        raise HTTPException(409, str(exc)) from exc

    await session.commit()
    return EditResolutionResponse(
        edit=DocumentEditRead.model_validate(updated),
        new_version=(
            DocumentVersionRead.model_validate(new_version) if new_version else None
        ),
        resolved_text=resolved_text,
    )


@router.post("/edits/{edit_id}/accept", response_model=EditResolutionResponse)
async def post_accept_edit(
    edit_id: uuid.UUID,
    session: AsyncSession = Depends(get_session),
    user: User = Depends(current_user),
) -> EditResolutionResponse:
    """Accept a single pending edit. Returns 409 if already resolved."""
    return await _resolve_one(None, edit_id, "accept", session, user)


@router.post("/edits/{edit_id}/reject", response_model=EditResolutionResponse)
async def post_reject_edit(
    edit_id: uuid.UUID,
    session: AsyncSession = Depends(get_session),
    user: User = Depends(current_user),
) -> EditResolutionResponse:
    """Reject a single pending edit. Returns 409 if already resolved."""
    return await _resolve_one(None, edit_id, "reject", session, user)


async def _resolve_all(
    version_id: uuid.UUID,
    action: Literal["accept_all", "reject_all"],
    session: AsyncSession,
    user: User,
) -> BulkResolutionResponse:
    try:
        affected, new_version, resolved_text = await resolve_bulk(
            session,
            version_id=version_id,
            actor_id=user.id,
            action=action,
        )
    except LookupError as exc:
        raise HTTPException(404, str(exc)) from exc
    await session.commit()
    return BulkResolutionResponse(
        affected_count=affected,
        new_version=DocumentVersionRead.model_validate(new_version),
        resolved_text=resolved_text,
    )


@router.post(
    "/versions/{version_id}/accept-all", response_model=BulkResolutionResponse
)
async def post_accept_all(
    version_id: uuid.UUID,
    session: AsyncSession = Depends(get_session),
    user: User = Depends(current_user),
) -> BulkResolutionResponse:
    """Accept every pending edit on this version in a single transaction."""
    return await _resolve_all(version_id, "accept_all", session, user)


@router.post(
    "/versions/{version_id}/reject-all", response_model=BulkResolutionResponse
)
async def post_reject_all(
    version_id: uuid.UUID,
    session: AsyncSession = Depends(get_session),
    user: User = Depends(current_user),
) -> BulkResolutionResponse:
    """Reject every pending edit on this version in a single transaction."""
    return await _resolve_all(version_id, "reject_all", session, user)


@router.get("/{document_id}/draft", response_model=DocumentWorkingDraftRead)
async def get_document_working_draft(
    document_id: uuid.UUID,
    session: AsyncSession = Depends(get_session),
    user: User = Depends(current_user),
) -> DocumentWorkingDraftRead:
    """Return the shared working draft for the document editor.

    If no mutable draft exists yet, return an initial draft derived from the
    latest saved text version or the extracted document body. This keeps
    initialisation server-owned without creating audit noise.
    """
    doc, _matter = await _load_owned_document(document_id, session, user)
    draft = await session.scalar(
        select(DocumentWorkingDraft).where(DocumentWorkingDraft.document_id == doc.id)
    )
    if draft is None:
        return await _initial_working_draft(session, doc)
    return DocumentWorkingDraftRead.model_validate(draft)


@router.put("/{document_id}/draft", response_model=DocumentWorkingDraftRead)
async def put_document_working_draft(
    document_id: uuid.UUID,
    body: DocumentWorkingDraftUpsert,
    session: AsyncSession = Depends(get_session),
    user: User = Depends(current_user),
) -> DocumentWorkingDraftRead:
    """Autosave the mutable working draft for a document.

    Draft writes are intentionally not audit rows. The matter record should
    show saved versions and professional decisions, not keystroke persistence.
    """
    doc, _matter = await _load_owned_document(document_id, session, user)
    base_version_id = body.base_version_id
    if base_version_id is not None:
        exists = await session.scalar(
            select(DocumentVersion.id).where(
                DocumentVersion.id == base_version_id,
                DocumentVersion.document_id == doc.id,
            )
        )
        if exists is None:
            raise HTTPException(
                422,
                {
                    "error": "invalid_base_version",
                    "message": "The draft base version does not belong to this document.",
                },
            )

    now = datetime.now(UTC)
    draft = await session.scalar(
        select(DocumentWorkingDraft).where(DocumentWorkingDraft.document_id == doc.id)
    )
    if draft is None:
        if body.expected_version_counter not in (None, 0):
            raise HTTPException(
                409,
                {
                    "error": "working_draft_conflict",
                    "message": "The shared draft changed before this save. Reload the document before saving again.",
                    "current_version_counter": 0,
                    "current_client_id": None,
                },
            )
        draft = DocumentWorkingDraft(
            document_id=doc.id,
            updated_by_id=user.id,
            updated_at=now,
            plain_text=body.plain_text,
            editor_json=body.editor_json,
            base_version_id=base_version_id,
            version_counter=1,
            client_id=body.client_id,
        )
        session.add(draft)
    else:
        if (
            body.expected_version_counter is not None
            and body.expected_version_counter != draft.version_counter
        ):
            raise HTTPException(
                409,
                {
                    "error": "working_draft_conflict",
                    "message": "The shared draft changed before this save. Reload the document before saving again.",
                    "current_version_counter": draft.version_counter,
                    "current_client_id": draft.client_id,
                },
            )
        draft.updated_by_id = user.id
        draft.updated_at = now
        draft.plain_text = body.plain_text
        draft.editor_json = body.editor_json
        draft.base_version_id = base_version_id
        draft.client_id = body.client_id
        draft.version_counter += 1

    await session.commit()
    return DocumentWorkingDraftRead.model_validate(draft)


@router.post("/{document_id}/draft/commit", response_model=DocumentVersionRead)
async def post_document_working_draft_commit(
    document_id: uuid.UUID,
    session: AsyncSession = Depends(get_session),
    user: User = Depends(current_user),
    body: DocumentWorkingDraftCommitRequest = DocumentWorkingDraftCommitRequest(),
) -> DocumentVersionRead:
    """Commit the mutable working draft as an immutable document version."""
    doc, matter = await _load_owned_document(document_id, session, user)
    draft = await session.scalar(
        select(DocumentWorkingDraft).where(DocumentWorkingDraft.document_id == doc.id)
    )
    if draft is None or not draft.plain_text.strip():
        raise HTTPException(
            422,
            {
                "error": "working_draft_empty",
                "message": "There is no working draft text to save as a version.",
            },
        )
    if (
        body.expected_version_counter is not None
        and body.expected_version_counter != draft.version_counter
    ):
        raise HTTPException(
            409,
            {
                "error": "working_draft_conflict",
                "message": "The shared draft changed before this version save. Reload the document before saving again.",
                "current_version_counter": draft.version_counter,
                "current_client_id": draft.client_id,
            },
        )

    version = await _create_user_edit_version(
        session,
        doc=doc,
        matter=matter,
        user=user,
        resolved_text=draft.plain_text,
        resolved_json=draft.editor_json,
        notes=body.notes or "Saved working draft from Legalise document editor",
        audit_source="working_draft",
    )
    if body.clear_draft:
        await session.delete(draft)
    await session.commit()
    return DocumentVersionRead.model_validate(version)


@router.post("/{document_id}/versions/manual", response_model=DocumentVersionRead)
async def post_manual_document_version(
    document_id: uuid.UUID,
    body: ManualDocumentVersionRequest,
    session: AsyncSession = Depends(get_session),
    user: User = Depends(current_user),
) -> DocumentVersionRead:
    """Save user-edited text as a new immutable document version."""
    doc, matter = await _load_owned_document(document_id, session, user)
    version = await _create_user_edit_version(
        session,
        doc=doc,
        matter=matter,
        user=user,
        resolved_text=body.resolved_text,
        resolved_json=body.resolved_json,
        notes=body.notes,
    )
    await session.commit()
    return DocumentVersionRead.model_validate(version)


@router.post("/{document_id}/versions/upload", response_model=DocumentVersionRead)
async def post_upload_document_version(
    document_id: uuid.UUID,
    file: UploadFile = File(...),
    notes: str | None = Form(default=None),
    session: AsyncSession = Depends(get_session),
    user: User = Depends(current_user),
) -> DocumentVersionRead:
    """Upload a replacement binary and make it the active document version.

    This is the binary counterpart to manual editor saves: the document keeps
    its identity, but the active original file, hash, extracted body, and
    version history move forward together.
    """
    doc, matter = await _load_owned_document(document_id, session, user)

    if file.content_type not in ALLOWED_UPLOAD_MIMES:
        raise HTTPException(
            415,
            detail={
                "error": "unsupported_mime",
                "got": file.content_type,
                "allowed": sorted(ALLOWED_UPLOAD_MIMES),
            },
        )

    contents = await file.read()
    if len(contents) > MAX_UPLOAD_BYTES:
        raise HTTPException(
            413,
            detail={
                "error": "upload_too_large",
                "max_bytes": MAX_UPLOAD_BYTES,
                "got_bytes": len(contents),
            },
        )

    declared_format = MIME_TO_FORMAT[file.content_type or ""]
    inferred_format = sniff_format(contents[:1024])
    if inferred_format is None or declared_format != inferred_format:
        raise HTTPException(
            415,
            detail={
                "error": "magic_byte_mismatch",
                "declared_mime": file.content_type,
                "declared_format": declared_format,
                "inferred_format": inferred_format,
            },
        )

    sha = hashlib.sha256(contents).hexdigest()
    filename = file.filename or doc.filename or "untitled"
    obj_key = uploaded_key(
        user_id=user.id,
        matter_id=matter.id,
        document_id=doc.id,
        sha256=sha,
    )
    storage = get_storage_backend()
    try:
        storage.put_bytes(
            obj_key,
            contents,
            content_type=file.content_type or "application/octet-stream",
            metadata={
                "filename": filename[:200],
                "sha256": sha,
                "document_id": str(doc.id),
            },
        )
    except StorageWriteError as exc:
        await audit_failure(
            session,
            "storage.put_bytes.failed",
            actor_id=user.id,
            matter_id=matter.id,
            module="storage",
            resource_type="document",
            resource_id=str(doc.id),
            payload={
                "storage_key": obj_key,
                "backend": exc.backend,
                "error_code": exc.error_code,
                "version_upload": True,
            },
        )
        raise HTTPException(
            502,
            detail={
                "error": "storage_write_failed",
                "message": "Failed to write document version to object storage.",
                "storage_key": obj_key,
                "backend": exc.backend,
            },
        ) from exc

    next_version = (
        await session.scalar(
            select(func.coalesce(func.max(DocumentVersion.version_number), 0) + 1)
            .where(DocumentVersion.document_id == doc.id)
        )
        or 1
    )
    extract_result = extract_text(
        contents,
        file.content_type or "application/octet-stream",
        filename,
    )
    version = DocumentVersion(
        document_id=doc.id,
        version_number=int(next_version),
        kind=VERSION_KIND_UPLOAD,
        created_by_id=user.id,
        storage_uri=obj_key,
        filename=filename,
        mime_type=file.content_type or "application/octet-stream",
        size_bytes=len(contents),
        sha256=sha,
        notes=notes or f"Uploaded replacement file: {filename}",
        resolved_text=(
            extract_result.extracted_text
            if extract_result.extraction_method != "failed"
            else None
        ),
    )
    session.add(version)

    doc.filename = filename
    doc.mime_type = file.content_type or "application/octet-stream"
    doc.size_bytes = len(contents)
    doc.sha256 = sha
    doc.storage_uri = obj_key
    doc.uploaded_at = datetime.now(UTC)
    doc.uploaded_by_id = user.id

    body = await session.scalar(
        select(DocumentBody).where(
            DocumentBody.document_id == doc.id,
            DocumentBody.kind == BODY_KIND_EXTRACTED,
        )
    )
    body_payload = {
        "extracted_text": extract_result.extracted_text,
        "extraction_method": extract_result.extraction_method,
        "char_count": extract_result.char_count,
        "page_count": extract_result.page_count,
        "error_reason": extract_result.error_reason,
        "extracted_at": datetime.now(UTC),
    }
    if body is None:
        session.add(
            DocumentBody(
                document_id=doc.id,
                kind=BODY_KIND_EXTRACTED,
                **body_payload,
            )
        )
    else:
        for key, value in body_payload.items():
            setattr(body, key, value)

    await session.flush()
    await audit.log(
        session,
        "document.version.uploaded",
        actor_id=user.id,
        matter_id=matter.id,
        module="document_editor",
        resource_type="document_version",
        resource_id=str(version.id),
        payload={
            "document_id": str(doc.id),
            "version_number": version.version_number,
            "filename": filename,
            "sha256": sha,
            "mime_type": doc.mime_type,
            "size_bytes": doc.size_bytes,
        },
    )
    await audit.log(
        session,
        (
            "document.text_extraction_failed"
            if extract_result.extraction_method == "failed"
            else "document.text_extracted"
        ),
        actor_id=user.id,
        matter_id=matter.id,
        module="document_ingestion",
        resource_type="document",
        resource_id=str(doc.id),
        payload={
            "version_id": str(version.id),
            "version_number": version.version_number,
            "method": extract_result.extraction_method,
            "char_count": extract_result.char_count,
            "page_count": extract_result.page_count,
            "mime_type": doc.mime_type,
            "reason": extract_result.error_reason,
        },
    )
    await session.commit()
    return DocumentVersionRead.model_validate(version)


@router.post("/{document_id}/versions/{version_id}/restore", response_model=DocumentVersionRead)
async def post_restore_document_version(
    document_id: uuid.UUID,
    version_id: uuid.UUID,
    body: RestoreDocumentVersionRequest,
    session: AsyncSession = Depends(get_session),
    user: User = Depends(current_user),
) -> DocumentVersionRead:
    """Restore a prior saved version as the active document.

    The original history remains immutable: restore creates a new
    `restored` version row and updates the active document pointer/body
    to match the selected version.
    """
    doc, matter = await _load_owned_document(document_id, session, user)
    source = await session.scalar(
        select(DocumentVersion).where(
            DocumentVersion.id == version_id,
            DocumentVersion.document_id == doc.id,
        )
    )
    if source is None:
        raise HTTPException(404, "document version not found")
    if not source.storage_uri and not source.resolved_text:
        raise HTTPException(422, "document version has no restorable content")

    filename = source.filename or doc.filename
    mime_type = source.mime_type or doc.mime_type or "application/octet-stream"
    storage_uri = source.storage_uri
    size_bytes = source.size_bytes
    sha = source.sha256
    resolved_text = source.resolved_text
    resolved_json = source.resolved_json
    extraction_method = "passthrough"
    page_count: int | None = None
    error_reason: str | None = None

    if storage_uri:
        storage = get_storage_backend()
        try:
            contents = storage.get_bytes(storage_uri)
        except KeyError:
            raise HTTPException(
                410,
                detail={
                    "error": "document_version_original_unavailable",
                    "message": "This saved version points to a file that is no longer available.",
                },
            )
        except StorageReadError as exc:
            await audit_failure(
                session,
                "storage.get_bytes.failed",
                actor_id=user.id,
                matter_id=matter.id,
                module="storage",
                resource_type="document_version",
                resource_id=str(source.id),
                payload={
                    "storage_key": storage_uri,
                    "backend": exc.backend,
                    "error_code": exc.error_code,
                    "restore": True,
                },
            )
            raise HTTPException(
                410,
                detail={
                    "error": "document_version_original_unavailable",
                    "message": "This saved version points to a file that is no longer available.",
                },
            ) from exc
        size_bytes = len(contents)
        sha = hashlib.sha256(contents).hexdigest()
        extracted = extract_text(contents, mime_type, filename)
        resolved_text = (
            extracted.extracted_text
            if extracted.extraction_method != "failed"
            else resolved_text
        )
        extraction_method = extracted.extraction_method
        page_count = extracted.page_count
        error_reason = extracted.error_reason
    else:
        contents = (resolved_text or "").encode("utf-8")
        sha = hashlib.sha256(contents).hexdigest()
        size_bytes = len(contents)
        mime_type = "text/plain"
        if not filename.lower().endswith(".txt"):
            filename = f"{filename.rsplit('.', 1)[0]}-v{source.version_number}.txt"
        storage_uri = uploaded_key(
            user_id=user.id,
            matter_id=matter.id,
            document_id=doc.id,
            sha256=sha,
        )
        try:
            get_storage_backend().put_bytes(
                storage_uri,
                contents,
                content_type=mime_type,
                metadata={
                    "filename": filename[:200],
                    "sha256": sha,
                    "document_id": str(doc.id),
                    "restored_from_version_id": str(source.id),
                },
            )
        except StorageWriteError as exc:
            await audit_failure(
                session,
                "storage.put_bytes.failed",
                actor_id=user.id,
                matter_id=matter.id,
                module="storage",
                resource_type="document_version",
                resource_id=str(source.id),
                payload={
                    "storage_key": storage_uri,
                    "backend": exc.backend,
                    "error_code": exc.error_code,
                    "restore": True,
                },
            )
            raise HTTPException(
                502,
                detail={
                    "error": "storage_write_failed",
                    "message": "Failed to store restored document version.",
                    "storage_key": storage_uri,
                    "backend": exc.backend,
                },
            ) from exc

    next_version = (
        await session.scalar(
            select(func.coalesce(func.max(DocumentVersion.version_number), 0) + 1)
            .where(DocumentVersion.document_id == doc.id)
        )
        or 1
    )
    restored = DocumentVersion(
        document_id=doc.id,
        version_number=int(next_version),
        kind=VERSION_KIND_RESTORED,
        created_by_id=user.id,
        storage_uri=storage_uri,
        filename=filename,
        mime_type=mime_type,
        size_bytes=size_bytes,
        sha256=sha,
        notes=body.notes or f"Restored from v{source.version_number}",
        resolved_text=resolved_text,
        resolved_json=resolved_json,
    )
    session.add(restored)

    doc.filename = filename
    doc.mime_type = mime_type
    doc.size_bytes = int(size_bytes or 0)
    doc.sha256 = sha or doc.sha256
    doc.storage_uri = storage_uri
    doc.uploaded_at = datetime.now(UTC)
    doc.uploaded_by_id = user.id

    active_body = await session.scalar(
        select(DocumentBody).where(
            DocumentBody.document_id == doc.id,
            DocumentBody.kind == BODY_KIND_EXTRACTED,
        )
    )
    body_payload = {
        "extracted_text": resolved_text or "",
        "extraction_method": extraction_method,
        "char_count": len(resolved_text or ""),
        "page_count": page_count,
        "error_reason": error_reason,
        "extracted_at": datetime.now(UTC),
    }
    if active_body is None:
        session.add(DocumentBody(document_id=doc.id, kind=BODY_KIND_EXTRACTED, **body_payload))
    else:
        for key, value in body_payload.items():
            setattr(active_body, key, value)

    await session.flush()
    await audit.log(
        session,
        "document.version.restored",
        actor_id=user.id,
        matter_id=matter.id,
        module="document_editor",
        resource_type="document_version",
        resource_id=str(restored.id),
        payload={
            "document_id": str(doc.id),
            "restored_version_number": restored.version_number,
            "source_version_id": str(source.id),
            "source_version_number": source.version_number,
            "filename": filename,
            "sha256": sha,
            "mime_type": mime_type,
            "size_bytes": size_bytes,
        },
    )
    await session.commit()
    return DocumentVersionRead.model_validate(restored)


@router.get("/{document_id}/versions/{version_id}/docx")
async def get_document_version_docx(
    document_id: uuid.UUID,
    version_id: uuid.UUID,
    session: AsyncSession = Depends(get_session),
    user: User = Depends(current_user),
) -> StreamingResponse:
    """Download a saved document version as a Word document."""
    doc, matter = await _load_owned_document(document_id, session, user)
    version = await session.scalar(
        select(DocumentVersion).where(
            DocumentVersion.id == version_id,
            DocumentVersion.document_id == doc.id,
        )
    )
    if version is None:
        raise HTTPException(404, "document version not found")
    if not version.resolved_text:
        raise HTTPException(422, "document version has no resolved text")

    comments = (
        await session.execute(
            select(DocumentComment)
            .where(DocumentComment.document_id == doc.id)
            .order_by(DocumentComment.created_at.asc(), DocumentComment.id.asc())
        )
    ).scalars().all()
    data = (
        _render_tiptap_docx(
            doc.filename,
            version.resolved_json,
            version.resolved_text,
            comments,
            DocumentAssetContext(user.id, matter.id, doc.id),
        )
        if version.resolved_json
        else _render_resolved_text_docx(doc.filename, version.resolved_text, comments)
    )
    filename = _docx_export_filename(doc.filename, version.version_number)
    await audit.log(
        session,
        "document.version.docx.exported",
        actor_id=user.id,
        matter_id=matter.id,
        module="document_editor",
        resource_type="document_version",
        resource_id=str(version.id),
        payload={
            "document_id": str(doc.id),
            "version_number": version.version_number,
            "char_count": len(version.resolved_text),
            "byte_count": len(data),
            "format": "docx",
            "rich_json": version.resolved_json is not None,
            "review_note_count": len(comments),
        },
    )
    await session.commit()
    return StreamingResponse(
        iter([data]),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f'attachment; filename="{filename}"',
            "Content-Length": str(len(data)),
        },
    )


@router.get("/{document_id}/versions/{version_id}/pdf")
async def get_document_version_pdf(
    document_id: uuid.UUID,
    version_id: uuid.UUID,
    session: AsyncSession = Depends(get_session),
    user: User = Depends(current_user),
) -> StreamingResponse:
    """Download a saved document version as a print-ready PDF."""
    doc, matter = await _load_owned_document(document_id, session, user)
    version = await session.scalar(
        select(DocumentVersion).where(
            DocumentVersion.id == version_id,
            DocumentVersion.document_id == doc.id,
        )
    )
    if version is None:
        raise HTTPException(404, "document version not found")
    if not version.resolved_text:
        raise HTTPException(422, "document version has no resolved text")

    comments = (
        await session.execute(
            select(DocumentComment)
            .where(DocumentComment.document_id == doc.id)
            .order_by(DocumentComment.created_at.asc(), DocumentComment.id.asc())
        )
    ).scalars().all()
    html_doc = _render_document_version_html(
        doc.filename,
        version,
        comments,
        DocumentAssetContext(user.id, matter.id, doc.id),
    )
    try:
        data = await _html_to_pdf(html_doc)
    except RuntimeError as exc:
        raise HTTPException(
            502,
            detail={
                "error": "pdf_export_failed",
                "message": str(exc),
            },
        ) from exc
    filename = _pdf_export_filename(doc.filename, version.version_number)
    await audit.log(
        session,
        "document.version.pdf.exported",
        actor_id=user.id,
        matter_id=matter.id,
        module="document_editor",
        resource_type="document_version",
        resource_id=str(version.id),
        payload={
            "document_id": str(doc.id),
            "version_number": version.version_number,
            "char_count": len(version.resolved_text),
            "byte_count": len(data),
            "format": "pdf",
            "rich_json": version.resolved_json is not None,
            "review_note_count": len(comments),
        },
    )
    await session.commit()
    return StreamingResponse(
        iter([data]),
        media_type="application/pdf",
        headers={
            "Content-Disposition": f'attachment; filename="{filename}"',
            "Content-Length": str(len(data)),
        },
    )


@router.get("/{document_id}/versions/{version_id}/original")
async def get_document_version_original(
    document_id: uuid.UUID,
    version_id: uuid.UUID,
    download: int = Query(0),
    session: AsyncSession = Depends(get_session),
    user: User = Depends(current_user),
) -> StreamingResponse:
    """Stream the original uploaded bytes for a saved upload version."""
    doc, matter = await _load_owned_document(document_id, session, user)
    version = await session.scalar(
        select(DocumentVersion).where(
            DocumentVersion.id == version_id,
            DocumentVersion.document_id == doc.id,
        )
    )
    if version is None:
        raise HTTPException(404, "document version not found")
    if not version.storage_uri:
        raise HTTPException(404, "version original file not available")

    storage = get_storage_backend()
    try:
        data = storage.get_bytes(version.storage_uri)
    except KeyError:
        raise HTTPException(404, "version original file not available")
    except StorageReadError as exc:
        await audit_failure(
            session,
            "storage.get_bytes.failed",
            actor_id=user.id,
            matter_id=matter.id,
            module="storage",
            resource_type="document_version",
            resource_id=str(version.id),
            payload={
                "document_id": str(doc.id),
                "storage_key": version.storage_uri,
                "backend": exc.backend,
                "error_code": exc.error_code,
            },
        )
        raise HTTPException(
            502,
            detail={
                "error": "storage_read_failed",
                "message": "Failed to read the version original from object storage.",
                "storage_key": version.storage_uri,
                "backend": exc.backend,
            },
        ) from exc

    is_download = bool(download)
    await audit.log(
        session,
        "document.version.original.accessed",
        actor_id=user.id,
        matter_id=matter.id,
        module="document_editor",
        resource_type="document_version",
        resource_id=str(version.id),
        payload={
            "document_id": str(doc.id),
            "version_number": version.version_number,
            "storage_key": version.storage_uri,
            "filename": version.filename or doc.filename,
            "sha256": version.sha256,
            "mime_type": version.mime_type or doc.mime_type,
            "size_bytes": version.size_bytes,
            "download": is_download,
        },
    )
    await session.commit()

    filename = _safe_filename(version.filename or doc.filename, str(version.id))
    disposition = "attachment" if is_download else "inline"
    return StreamingResponse(
        iter([data]),
        media_type=version.mime_type or doc.mime_type or "application/octet-stream",
        headers={
            "Content-Disposition": f'{disposition}; filename="{filename}"',
            "Content-Length": str(len(data)),
        },
    )


@router.get("/{document_id}/versions", response_model=list[DocumentVersionSummary])
async def get_document_versions(
    document_id: uuid.UUID,
    session: AsyncSession = Depends(get_session),
    user: User = Depends(current_user),
) -> list[DocumentVersionSummary]:
    """List versions for a document with per-version edit counts.

    Returns versions ordered by `version_number` ascending. 404 if the
    document isn't owned by the current user.
    """
    pair = (
        await session.execute(
            select(Document, Matter)
            .join(Matter, Matter.id == Document.matter_id)
            .where(Document.id == document_id)
        )
    ).first()
    if pair is None:
        raise HTTPException(404, "document not found")
    _, matter = pair
    if matter.created_by_id != user.id or matter.status == STATUS_ARCHIVED:
        raise HTTPException(404, "document not found")

    versions = (
        await session.execute(
            select(DocumentVersion)
            .where(DocumentVersion.document_id == document_id)
            .order_by(DocumentVersion.version_number.asc())
        )
    ).scalars().all()

    if not versions:
        return []

    counts_rows = (
        await session.execute(
            select(
                DocumentEdit.document_version_id,
                DocumentEdit.status,
                func.count(DocumentEdit.id),
            )
            .where(
                DocumentEdit.document_version_id.in_([v.id for v in versions])
            )
            .group_by(DocumentEdit.document_version_id, DocumentEdit.status)
        )
    ).all()
    counts: dict[uuid.UUID, dict[str, int]] = {}
    for vid, status, n in counts_rows:
        counts.setdefault(vid, {})[status] = int(n)

    out: list[DocumentVersionSummary] = []
    for v in versions:
        c = counts.get(v.id, {})
        out.append(
            DocumentVersionSummary(
                version=DocumentVersionRead.model_validate(v),
                pending_count=c.get(EDIT_STATUS_PENDING, 0),
                accepted_count=c.get(EDIT_STATUS_ACCEPTED, 0),
                rejected_count=c.get(EDIT_STATUS_REJECTED, 0),
            )
        )
    return out


@router.get("/{document_id}/edit-sessions", response_model=list[DocumentEditSessionRead])
async def get_document_edit_sessions(
    document_id: uuid.UUID,
    session: AsyncSession = Depends(get_session),
    user: User = Depends(current_user),
) -> list[DocumentEditSessionRead]:
    """List active editing sessions for this document."""
    doc, _matter = await _load_owned_document(document_id, session, user)
    return await _active_edit_sessions(session, doc.id)


@router.post("/{document_id}/edit-sessions", response_model=DocumentEditSessionResponse)
async def post_document_edit_session(
    document_id: uuid.UUID,
    body: DocumentEditSessionStart,
    session: AsyncSession = Depends(get_session),
    user: User = Depends(current_user),
) -> DocumentEditSessionResponse:
    """Start or heartbeat a document editing session for this browser client."""
    doc, matter = await _load_owned_document(document_id, session, user)
    now = datetime.now(UTC)
    row = await session.scalar(
        select(DocumentEditSession).where(
            DocumentEditSession.document_id == doc.id,
            DocumentEditSession.user_id == user.id,
            DocumentEditSession.client_id == body.client_id,
            DocumentEditSession.ended_at.is_(None),
        )
    )
    started = False
    if row is None:
        row = DocumentEditSession(
            document_id=doc.id,
            user_id=user.id,
            client_id=body.client_id,
            started_at=now,
            last_seen_at=now,
        )
        session.add(row)
        started = True
    else:
        row.last_seen_at = now

    await session.flush()
    if started:
        await audit.log(
            session,
            "document.edit_session.started",
            actor_id=user.id,
            matter_id=matter.id,
            module="document_editor",
            resource_type="document",
            resource_id=str(doc.id),
            payload={"session_id": str(row.id)},
        )
    await session.commit()
    active = await _active_edit_sessions(session, doc.id)
    current = next((item for item in active if item.id == row.id), None)
    if current is None:
        current = DocumentEditSessionRead(
            id=row.id,
            document_id=row.document_id,
            user_id=row.user_id,
            client_id=row.client_id,
            user_label=user.name or user.email,
            started_at=row.started_at,
            last_seen_at=row.last_seen_at,
            ended_at=row.ended_at,
        )
    return DocumentEditSessionResponse(current=current, active=active)


@router.delete("/{document_id}/edit-sessions/{session_id}", status_code=204)
async def delete_document_edit_session(
    document_id: uuid.UUID,
    session_id: uuid.UUID,
    session: AsyncSession = Depends(get_session),
    user: User = Depends(current_user),
) -> Response:
    """End this browser client's document editing session."""
    doc, matter = await _load_owned_document(document_id, session, user)
    row = await session.scalar(
        select(DocumentEditSession).where(
            DocumentEditSession.id == session_id,
            DocumentEditSession.document_id == doc.id,
            DocumentEditSession.user_id == user.id,
        )
    )
    if row is not None and row.ended_at is None:
        row.ended_at = datetime.now(UTC)
        row.last_seen_at = row.ended_at
        await audit.log(
            session,
            "document.edit_session.ended",
            actor_id=user.id,
            matter_id=matter.id,
            module="document_editor",
            resource_type="document",
            resource_id=str(doc.id),
            payload={"session_id": str(row.id)},
        )
        await session.commit()
    return Response(status_code=204)


# -- Anonymisation ---------------------------------------------------------


async def _load_owned_document(
    document_id: uuid.UUID, session: AsyncSession, user: User
) -> tuple[Document, Matter]:
    """Resolve a document the current user owns, else 404.

    Walks document → matter → ownership in one round trip. Shared by
    every anonymisation endpoint so the auth shape doesn't drift.
    """
    row = (
        await session.execute(
            select(Document, Matter)
            .join(Matter, Matter.id == Document.matter_id)
            .where(Document.id == document_id)
        )
    ).first()
    if row is None:
        raise HTTPException(404, "document not found")
    doc, matter = row
    if matter.created_by_id != user.id or matter.status == STATUS_ARCHIVED:
        raise HTTPException(404, "document not found")
    return doc, matter


@router.get("/{document_id}/comments", response_model=list[DocumentCommentRead])
async def get_document_comments(
    document_id: uuid.UUID,
    session: AsyncSession = Depends(get_session),
    user: User = Depends(current_user),
) -> list[DocumentCommentRead]:
    """List review notes for an owned, live document."""
    doc, _matter = await _load_owned_document(document_id, session, user)
    comments = (
        await session.scalars(
            select(DocumentComment)
            .where(DocumentComment.document_id == doc.id)
            .order_by(DocumentComment.created_at.asc(), DocumentComment.id.asc())
        )
    ).all()
    return [DocumentCommentRead.model_validate(comment) for comment in comments]


@router.post("/{document_id}/comments", response_model=DocumentCommentRead)
async def post_document_comment(
    document_id: uuid.UUID,
    body: DocumentCommentCreate,
    session: AsyncSession = Depends(get_session),
    user: User = Depends(current_user),
) -> DocumentCommentRead:
    """Add a human review note to an owned document."""
    doc, matter = await _load_owned_document(document_id, session, user)
    comment_body = body.body.strip()
    if len(comment_body) < 2:
        raise HTTPException(422, "comment body is required")
    quote_text = body.quote_text.strip() if body.quote_text else None
    anchor_start = body.anchor_start
    anchor_end = body.anchor_end
    body_sha256: str | None = None
    has_anchor = anchor_start is not None or anchor_end is not None
    if has_anchor:
        if anchor_start is None or anchor_end is None or anchor_end <= anchor_start:
            raise HTTPException(
                422,
                {
                    "error": "invalid_comment_anchor",
                    "message": "Comment anchors require start and end positions.",
                },
            )
        extracted = await extracted_body_for(session, doc.id)
        if extracted is None:
            raise HTTPException(
                422,
                {
                    "error": "document_body_unavailable",
                    "message": "The document text is not available for anchoring.",
                },
            )
        source_text = extracted.extracted_text
        body_sha256 = hashlib.sha256(source_text.encode("utf-8")).hexdigest()
        if body.body_sha256 and body.body_sha256 != body_sha256:
            raise HTTPException(
                409,
                {
                    "error": "document_body_changed",
                    "message": "The document text changed before the note was saved.",
                },
            )
        if anchor_end > len(source_text):
            raise HTTPException(
                422,
                {
                    "error": "invalid_comment_anchor",
                    "message": "The selected range is outside the document text.",
                },
            )
        anchored_quote = source_text[anchor_start:anchor_end].strip()
        if not quote_text and anchored_quote:
            quote_text = anchored_quote[:2000]
    comment = DocumentComment(
        document_id=doc.id,
        author_id=user.id,
        quote_text=quote_text or None,
        body_sha256=body_sha256,
        anchor_start=anchor_start if has_anchor else None,
        anchor_end=anchor_end if has_anchor else None,
        body=comment_body,
        status=COMMENT_STATUS_OPEN,
    )
    session.add(comment)
    await session.flush()
    await audit.log(
        session,
        "document.comment.created",
        actor_id=user.id,
        matter_id=matter.id,
        module="document_editor",
        resource_type="document_comment",
        resource_id=str(comment.id),
        payload={
            "document_id": str(doc.id),
            "has_quote": bool(comment.quote_text),
            "has_anchor": has_anchor,
            "body_sha256": comment.body_sha256,
            "anchor_start": comment.anchor_start,
            "anchor_end": comment.anchor_end,
        },
    )
    await session.commit()
    return DocumentCommentRead.model_validate(comment)


@router.patch(
    "/{document_id}/comments/{comment_id}",
    response_model=DocumentCommentRead,
)
async def patch_document_comment(
    document_id: uuid.UUID,
    comment_id: uuid.UUID,
    body: DocumentCommentUpdate,
    session: AsyncSession = Depends(get_session),
    user: User = Depends(current_user),
) -> DocumentCommentRead:
    """Update the body of an open document review note."""
    doc, matter = await _load_owned_document(document_id, session, user)
    comment = await session.scalar(
        select(DocumentComment).where(
            DocumentComment.id == comment_id,
            DocumentComment.document_id == doc.id,
        )
    )
    if comment is None:
        raise HTTPException(404, "document comment not found")
    if comment.status == COMMENT_STATUS_RESOLVED:
        raise HTTPException(
            409,
            {
                "error": "document_comment_resolved",
                "message": "Resolved review notes cannot be edited.",
            },
        )
    next_body = body.body.strip()
    if len(next_body) < 2:
        raise HTTPException(422, "comment body is required")
    if next_body != comment.body:
        previous_length = len(comment.body)
        comment.body = next_body
        await audit.log(
            session,
            "document.comment.updated",
            actor_id=user.id,
            matter_id=matter.id,
            module="document_editor",
            resource_type="document_comment",
            resource_id=str(comment.id),
            payload={
                "document_id": str(doc.id),
                "previous_length": previous_length,
                "next_length": len(next_body),
            },
        )
    await session.commit()
    return DocumentCommentRead.model_validate(comment)


@router.post(
    "/{document_id}/comments/{comment_id}/resolve",
    response_model=DocumentCommentRead,
)
async def post_resolve_document_comment(
    document_id: uuid.UUID,
    comment_id: uuid.UUID,
    session: AsyncSession = Depends(get_session),
    user: User = Depends(current_user),
) -> DocumentCommentRead:
    """Resolve a document review note."""
    doc, matter = await _load_owned_document(document_id, session, user)
    comment = await session.scalar(
        select(DocumentComment).where(
            DocumentComment.id == comment_id,
            DocumentComment.document_id == doc.id,
        )
    )
    if comment is None:
        raise HTTPException(404, "document comment not found")
    if comment.status != COMMENT_STATUS_RESOLVED:
        comment.status = COMMENT_STATUS_RESOLVED
        comment.resolved_at = datetime.now(UTC)
        comment.resolved_by_id = user.id
        await audit.log(
            session,
            "document.comment.resolved",
            actor_id=user.id,
            matter_id=matter.id,
            module="document_editor",
            resource_type="document_comment",
            resource_id=str(comment.id),
            payload={"document_id": str(doc.id)},
        )
    await session.commit()
    return DocumentCommentRead.model_validate(comment)


@router.post(
    "/{document_id}/comments/{comment_id}/reopen",
    response_model=DocumentCommentRead,
)
async def post_reopen_document_comment(
    document_id: uuid.UUID,
    comment_id: uuid.UUID,
    session: AsyncSession = Depends(get_session),
    user: User = Depends(current_user),
) -> DocumentCommentRead:
    """Reopen a resolved document review note."""
    doc, matter = await _load_owned_document(document_id, session, user)
    comment = await session.scalar(
        select(DocumentComment).where(
            DocumentComment.id == comment_id,
            DocumentComment.document_id == doc.id,
        )
    )
    if comment is None:
        raise HTTPException(404, "document comment not found")
    if comment.status == COMMENT_STATUS_RESOLVED:
        comment.status = COMMENT_STATUS_OPEN
        comment.resolved_at = None
        comment.resolved_by_id = None
        await audit.log(
            session,
            "document.comment.reopened",
            actor_id=user.id,
            matter_id=matter.id,
            module="document_editor",
            resource_type="document_comment",
            resource_id=str(comment.id),
            payload={"document_id": str(doc.id)},
        )
    await session.commit()
    return DocumentCommentRead.model_validate(comment)


def _result_from_redacted(body: DocumentBody) -> AnonymisationResult:
    """Reconstruct an AnonymisationResult from the persisted redacted body."""
    tokens: list[TokenMapping] = []
    mapping = body.mapping if isinstance(body.mapping, dict) else {}
    token_map = mapping.get("tokens") if isinstance(mapping, dict) else None
    if isinstance(token_map, dict):
        for token, meta in token_map.items():
            if not isinstance(token, str) or not isinstance(meta, dict):
                continue
            tokens.append(
                TokenMapping(
                    token=token,
                    entity_type=str(meta.get("entity_type", "")),
                    original=str(meta.get("original", "")),
                    occurrences=int(meta.get("occurrences", 0) or 0),
                )
            )
    tokens.sort(key=lambda t: (t.token.split("_", 1)[0], t.token))
    return AnonymisationResult(
        document_id=body.document_id,
        redacted_text=body.extracted_text,
        engine=body.engine or body.extraction_method or "unknown",
        anonymised_at=body.anonymised_at or body.extracted_at,
        char_count=body.char_count,
        entity_count=len(tokens),
        tokens=tokens,
    )


@router.post("/{document_id}/anonymise", response_model=AnonymisationResult)
async def post_anonymise_document(
    document_id: uuid.UUID,
    body: AnonymiseRequest,
    session: AsyncSession = Depends(get_session),
    user: User = Depends(current_user),
) -> AnonymisationResult:
    """Run anonymisation and UPSERT the `redacted` DocumentBody."""
    try:
        result = await anonymise_document(
            session=session,
            gateway=model_gateway,
            document_id=document_id,
            actor_id=user.id,
            engine=body.engine,
            entity_types=body.entity_types,
            threshold=body.threshold,
        )
    except LookupError as exc:
        raise HTTPException(404, str(exc)) from exc
    except ValueError as exc:
        raise HTTPException(422, str(exc)) from exc
    except PROVIDER_HTTP_EXCEPTIONS as exc:
        raise provider_error_http_exception(exc) from exc
    except RuntimeError as exc:
        # Presidio not installed in this environment. 503 communicates
        # "service is real but disabled" more honestly than 500.
        raise HTTPException(503, f"anonymisation engine unavailable: {exc}") from exc

    await session.commit()
    return result


@router.get("/{document_id}/anonymise", response_model=AnonymisationResult)
async def get_anonymise_document(
    document_id: uuid.UUID,
    session: AsyncSession = Depends(get_session),
    user: User = Depends(current_user),
) -> AnonymisationResult:
    """Return the most recent redacted body for this document."""
    await _load_owned_document(document_id, session, user)
    redacted = await session.scalar(
        select(DocumentBody).where(
            DocumentBody.document_id == document_id,
            DocumentBody.kind == BODY_KIND_REDACTED,
        )
    )
    if redacted is None:
        raise HTTPException(404, "no anonymised body for this document")
    return _result_from_redacted(redacted)


@router.get("/{document_id}/anonymise/mapping", response_model=MappingRead)
async def get_anonymise_mapping(
    document_id: uuid.UUID,
    session: AsyncSession = Depends(get_session),
    user: User = Depends(current_user),
) -> MappingRead:
    """Return the token → original mapping. Matter-owner-only.

    `_load_owned_document` already enforces owner-only via 404; we
    additionally write a `module.anonymisation.viewed` audit row so
    mapping reveals are traceable.
    """
    doc, matter = await _load_owned_document(document_id, session, user)
    redacted = await session.scalar(
        select(DocumentBody).where(
            DocumentBody.document_id == doc.id,
            DocumentBody.kind == BODY_KIND_REDACTED,
        )
    )
    if redacted is None:
        raise HTTPException(404, "no anonymised body for this document")

    mapping = redacted.mapping if isinstance(redacted.mapping, dict) else {}
    token_map = mapping.get("tokens") if isinstance(mapping, dict) else None
    tokens: list[TokenMapping] = []
    if isinstance(token_map, dict):
        for token, meta in token_map.items():
            if not isinstance(token, str) or not isinstance(meta, dict):
                continue
            tokens.append(
                TokenMapping(
                    token=token,
                    entity_type=str(meta.get("entity_type", "")),
                    original=str(meta.get("original", "")),
                    occurrences=int(meta.get("occurrences", 0) or 0),
                )
            )
    tokens.sort(key=lambda t: (t.token.split("_", 1)[0], t.token))

    raw_spans = mapping.get("spans") if isinstance(mapping, dict) else None
    spans = raw_spans if isinstance(raw_spans, list) else []

    await audit.log(
        session,
        "module.anonymisation.viewed",
        actor_id=user.id,
        matter_id=matter.id,
        module="anonymisation",
        resource_type="document",
        resource_id=str(doc.id),
        payload={"token_count": len(tokens)},
    )
    await session.commit()

    return MappingRead(document_id=doc.id, tokens=tokens, spans=spans)


@router.delete(
    "/{document_id}/anonymise",
    status_code=204,
    response_class=Response,
)
async def delete_anonymise_document(
    document_id: uuid.UUID,
    session: AsyncSession = Depends(get_session),
    user: User = Depends(current_user),
) -> Response:
    """Delete the redacted DocumentBody so the next run starts cold."""
    doc, matter = await _load_owned_document(document_id, session, user)
    redacted = await session.scalar(
        select(DocumentBody).where(
            DocumentBody.document_id == doc.id,
            DocumentBody.kind == BODY_KIND_REDACTED,
        )
    )
    if redacted is None:
        return Response(status_code=204)
    await session.delete(redacted)
    await audit.log(
        session,
        "module.anonymisation.deleted",
        actor_id=user.id,
        matter_id=matter.id,
        module="anonymisation",
        resource_type="document",
        resource_id=str(doc.id),
        payload={"engine": redacted.engine},
    )
    await session.commit()
    return Response(status_code=204)
