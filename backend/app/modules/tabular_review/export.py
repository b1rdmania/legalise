"""Landscape .docx export for a tabular review.

Renders directly via python-docx (a real `doc.add_table` table) rather
than routing through the markdown-based `generate_docx` tool. Cleaner
module separation: this export depends on no other workstream's
behaviour, and the docx-table feature lives next to its single caller.

Filename convention: `tabular-review-{slug}-{title-slug}.docx`.
Storage path mirrors the `generate_docx` convention so the existing
`GET /api/documents/generated/{file_uuid}` endpoint can stream the file
without modification — it looks up `document.generated` audit rows by
`resource_id=str(file_uuid)`, so we emit one with `format=docx`.
"""

from __future__ import annotations

import re
import uuid
from datetime import datetime, timezone
from io import BytesIO
from typing import Sequence

from docx import Document as DocxDocument
from docx.enum.section import WD_ORIENTATION
from docx.shared import Pt
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.api import audit as audit_api
from app.core.storage import generated_key, get_storage_backend
from app.models import Document, Matter
from app.models.tabular_review import TabularReview, TabularReviewRow

from .schemas import ColumnSpec


_FILENAME_SAFE_RE = re.compile(r"[^A-Za-z0-9._-]+")


def _slug_title(title: str) -> str:
    cleaned = _FILENAME_SAFE_RE.sub("-", title).strip("-._").lower()
    return cleaned[:64] or "review"


async def _load_grid(
    session: AsyncSession, matter_id: uuid.UUID, review_id: uuid.UUID
) -> tuple[list[Document], dict[uuid.UUID, dict]]:
    documents = list(
        (
            await session.scalars(
                select(Document)
                .where(Document.matter_id == matter_id)
                .order_by(Document.uploaded_at.asc())
            )
        ).all()
    )
    rows = list(
        (
            await session.scalars(
                select(TabularReviewRow).where(TabularReviewRow.review_id == review_id)
            )
        ).all()
    )
    values_by_doc: dict[uuid.UUID, dict] = {
        r.document_id: dict(r.extracted_values or {}) for r in rows
    }
    return documents, values_by_doc


def _apply_landscape(document: DocxDocument) -> None:
    section = document.sections[0]
    new_width, new_height = section.page_height, section.page_width
    section.orientation = WD_ORIENTATION.LANDSCAPE
    section.page_width = new_width
    section.page_height = new_height


def _render_table(
    document: DocxDocument,
    columns: Sequence[ColumnSpec],
    documents: Sequence[Document],
    values_by_doc: dict[uuid.UUID, dict],
) -> None:
    headers = ["Filename"] + [c.label for c in columns]
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)

    for doc in documents:
        row_cells = table.add_row().cells
        row_cells[0].text = doc.filename
        values = values_by_doc.get(doc.id, {})
        for i, col in enumerate(columns, start=1):
            raw = values.get(col.key)
            row_cells[i].text = "" if raw is None else str(raw)
            for p in row_cells[i].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)


async def export_review_docx(
    *,
    session: AsyncSession,
    review: TabularReview,
    matter: Matter,
    actor_id: uuid.UUID,
) -> tuple[uuid.UUID, int, str]:
    """Render and persist a landscape .docx of `review`.

    Returns `(file_uuid, byte_count, storage_uri)`. Caller is responsible
    for the session commit so the audit row lands in the same transaction.
    """
    columns = [ColumnSpec.model_validate(c) for c in (review.columns_config or [])]
    documents, values_by_doc = await _load_grid(session, matter.id, review.id)

    docx = DocxDocument()
    _apply_landscape(docx)
    docx.add_heading(review.title, level=0)
    docx.add_paragraph(f"Matter: {matter.title} ({matter.slug})")
    docx.add_paragraph(
        f"Exported: {datetime.now(timezone.utc).isoformat(timespec='seconds')}"
    )
    if columns:
        _render_table(docx, columns, documents, values_by_doc)
    else:
        docx.add_paragraph("(No columns defined — nothing to export.)")

    file_uuid = uuid.uuid4()
    filename = f"{file_uuid}.docx"

    buf = BytesIO()
    docx.save(buf)
    payload_bytes = buf.getvalue()
    byte_count = len(payload_bytes)

    # Write to object storage (Unit 1 abstraction). The previous
    # filesystem write at settings.matters_root left bytes inaccessible
    # to the documents.py download path (which reads via storage in S3
    # production). Per HANDOVER_SUBSTRATE_REVIEW_FIXES.md §2 P1.
    storage_uri = generated_key(
        user_id=matter.created_by_id,
        matter_id=matter.id,
        document_id=file_uuid,
        filename=filename,
    )
    get_storage_backend().put_bytes(
        storage_uri,
        payload_bytes,
        content_type=(
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ),
        metadata={
            "matter_id": str(matter.id),
            "review_id": str(review.id),
            "actor_id": str(actor_id),
        },
    )

    title = f"tabular-review-{matter.slug}-{_slug_title(review.title)}"

    # Two audit rows: the cross-module `document.generated` (matches the
    # download endpoint's lookup) + the module-namespaced `exported` row.
    await audit_api.log(
        session,
        "document.generated",
        module="document_generation",
        actor_id=actor_id,
        matter_id=matter.id,
        resource_type="document",
        resource_id=str(file_uuid),
        payload={
            "format": "docx",
            "byte_count": byte_count,
            "storage_uri": storage_uri,
            "title": title,
            "source_module": "tabular_review",
        },
    )
    await audit_api.log(
        session,
        "module.tabular_review.exported",
        module="tabular_review",
        actor_id=actor_id,
        matter_id=matter.id,
        resource_type="document",
        resource_id=str(file_uuid),
        payload={
            "format": "docx",
            "review_id": str(review.id),
            "byte_count": byte_count,
            "storage_uri": storage_uri,
            "title": title,
        },
    )

    return file_uuid, byte_count, storage_uri


__all__ = ["export_review_docx"]
