"""Document engine contract tests."""

from __future__ import annotations

import io

from docx import Document as DocxDocument

from app.core.document_engine import blocks_from_docx, blocks_from_text


def test_blocks_from_text_ignores_empty_paragraphs() -> None:
    blocks = blocks_from_text("First paragraph.\n\n\nSecond paragraph.\n")

    assert [block.id for block in blocks] == ["p1", "p2"]
    assert [block.text for block in blocks] == [
        "First paragraph.",
        "Second paragraph.",
    ]
    assert all(block.type == "paragraph" for block in blocks)


def test_blocks_from_docx_extracts_paragraphs_and_table_cells() -> None:
    document = DocxDocument()
    document.add_paragraph("Main clause")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Party A"
    table.cell(0, 1).text = "Party B"
    buf = io.BytesIO()
    document.save(buf)

    blocks = blocks_from_docx(buf.getvalue())

    assert [block.text for block in blocks] == ["Main clause", "Party A", "Party B"]
    assert [block.type for block in blocks] == [
        "paragraph",
        "table_cell",
        "table_cell",
    ]
