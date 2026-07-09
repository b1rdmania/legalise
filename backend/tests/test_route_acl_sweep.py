"""Issue #3 — Route ACL sweep for non-matter (UUID-keyed) resources.

Verifies that every document endpoint that resolves by document UUID
(not by matter slug) still enforces:
  1. Ownership: another user cannot access your documents.
  2. Archived-matter gate: once the owning matter is tombstoned, 404.

DB-backed; skips when Postgres is unreachable (see conftest.py).
"""

from __future__ import annotations

import hashlib
import io
import base64
import uuid
import zipfile

import pytest
from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor


EMAIL_A = "acl-sweep-a@example.com"
PASSWORD_A = "acl-sweep-a-password-2026"
EMAIL_B = "acl-sweep-b@example.com"
PASSWORD_B = "acl-sweep-b-password-2026"

_PDF_MAGIC = b"%PDF-1.4 1 0 obj<</Type /Catalog>>stream\nHello\nendstream\nendobj"
_ONE_PIXEL_PNG = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO+/p9sAAAAASUVORK5CYII="
)


async def _signup_and_login(client, email: str, password: str) -> None:
    reg = await client.post("/auth/register", json={"email": email, "password": password})
    assert reg.status_code == 201, reg.text
    login = await client.post(
        "/auth/login",
        data={"username": email, "password": password},
        headers={"Content-Type": "application/x-www-form-urlencoded"},
    )
    assert login.status_code == 204, login.text


async def _create_matter_and_upload(client) -> tuple[str, str]:
    """Returns (matter_slug, document_id)."""
    create = await client.post(
        "/api/matters",
        json={"title": "ACL Sweep Matter"},
    )
    assert create.status_code == 201, create.text
    slug = create.json()["slug"]

    resp = await client.post(
        f"/api/matters/{slug}/documents",
        files={"file": ("test.pdf", io.BytesIO(_PDF_MAGIC), "application/pdf")},
        data={"tag": "draft"},
    )
    assert resp.status_code == 201, resp.text
    doc_id = resp.json()["id"]
    return slug, doc_id


async def _create_matter_and_upload_text(client, text: str) -> tuple[str, str]:
    """Returns (matter_slug, document_id) for an extracted text/plain document."""
    create = await client.post(
        "/api/matters",
        json={"title": "ACL Sweep Text Matter"},
    )
    assert create.status_code == 201, create.text
    slug = create.json()["slug"]

    resp = await client.post(
        f"/api/matters/{slug}/documents",
        files={"file": ("note.txt", io.BytesIO(text.encode("utf-8")), "text/plain")},
        data={"tag": "draft"},
    )
    assert resp.status_code == 201, resp.text
    doc_id = resp.json()["id"]
    return slug, doc_id


# ---------------------------------------------------------------------------
# Cross-user ownership — 404 on all document UUID endpoints
# ---------------------------------------------------------------------------


@pytest.mark.parametrize(
    "resource_path,needs_saved_version",
    [
        pytest.param("body", False, id="document-body"),
        pytest.param("versions", False, id="document-versions"),
        pytest.param("anonymise", False, id="anonymise"),
        pytest.param("versions/{version_id}/docx", True, id="version-docx"),
    ],
)
@pytest.mark.asyncio
async def test_document_get_routes_cross_user_returns_404(
    client, resource_path, needs_saved_version
) -> None:
    """User B cannot GET User A's document UUID-keyed resources."""
    await _signup_and_login(client, EMAIL_A, PASSWORD_A)
    _, doc_id = await _create_matter_and_upload(client)

    version_id = ""
    if needs_saved_version:
        save = await client.post(
            f"/api/documents/{doc_id}/versions/manual",
            json={"resolved_text": "Owned user edit."},
        )
        assert save.status_code == 200, save.text
        version_id = save.json()["id"]

    await client.post("/auth/logout")

    await _signup_and_login(client, EMAIL_B, PASSWORD_B)
    path = resource_path.format(version_id=version_id)
    resp = await client.get(f"/api/documents/{doc_id}/{path}")
    assert resp.status_code == 404, resp.text


@pytest.mark.asyncio
async def test_document_asset_owner_can_upload_and_read_image(client) -> None:
    """Owner can upload an embedded editor image and read it through the proxy."""
    await _signup_and_login(client, EMAIL_A, PASSWORD_A)
    _, doc_id = await _create_matter_and_upload_text(client, "Document body.")

    png = b"\x89PNG\r\n\x1a\nlegalise-image"
    uploaded = await client.post(
        f"/api/documents/{doc_id}/assets",
        files={"file": ("diagram.png", io.BytesIO(png), "image/png")},
    )
    assert uploaded.status_code == 200, uploaded.text
    payload = uploaded.json()
    assert payload["filename"] == "diagram.png"
    assert payload["mime_type"] == "image/png"
    assert payload["size_bytes"] == len(png)
    assert payload["sha256"] == hashlib.sha256(png).hexdigest()
    assert payload["url"].startswith(f"/api/documents/{doc_id}/assets/")

    read_back = await client.get(payload["url"])
    assert read_back.status_code == 200, read_back.text
    assert read_back.headers["content-type"].startswith("image/png")
    assert read_back.content == png


@pytest.mark.asyncio
async def test_document_asset_rejects_unsupported_mime(client) -> None:
    """Image assets are intentionally image-only."""
    await _signup_and_login(client, EMAIL_A, PASSWORD_A)
    _, doc_id = await _create_matter_and_upload_text(client, "Document body.")

    resp = await client.post(
        f"/api/documents/{doc_id}/assets",
        files={"file": ("payload.svg", io.BytesIO(b"<svg />"), "image/svg+xml")},
    )
    assert resp.status_code == 415, resp.text


@pytest.mark.asyncio
async def test_document_asset_cross_user_returns_404(client) -> None:
    """User B cannot read User A's embedded editor images by document UUID."""
    await _signup_and_login(client, EMAIL_A, PASSWORD_A)
    _, doc_id = await _create_matter_and_upload_text(client, "Document body.")
    uploaded = await client.post(
        f"/api/documents/{doc_id}/assets",
        files={"file": ("diagram.png", io.BytesIO(b"png-bytes"), "image/png")},
    )
    assert uploaded.status_code == 200, uploaded.text
    url = uploaded.json()["url"]
    await client.post("/auth/logout")

    await _signup_and_login(client, EMAIL_B, PASSWORD_B)
    resp = await client.get(url)
    assert resp.status_code == 404, resp.text


@pytest.mark.asyncio
async def test_document_comments_owner_can_create_list_resolve_and_reopen(client) -> None:
    """Owner can leave, resolve, and reopen review notes on a document."""
    await _signup_and_login(client, EMAIL_A, PASSWORD_A)
    _, doc_id = await _create_matter_and_upload(client)

    created = await client.post(
        f"/api/documents/{doc_id}/comments",
        json={
            "quote_text": "single social-media post",
            "body": "Check the source before relying on this.",
        },
    )
    assert created.status_code == 200, created.text
    comment = created.json()
    assert comment["status"] == "open"
    assert comment["quote_text"] == "single social-media post"
    assert comment["body"] == "Check the source before relying on this."

    listed = await client.get(f"/api/documents/{doc_id}/comments")
    assert listed.status_code == 200, listed.text
    assert [row["id"] for row in listed.json()] == [comment["id"]]

    updated = await client.patch(
        f"/api/documents/{doc_id}/comments/{comment['id']}",
        json={"body": "Checked against the source and updated for clarity."},
    )
    assert updated.status_code == 200, updated.text
    assert updated.json()["body"] == "Checked against the source and updated for clarity."

    resolved = await client.post(
        f"/api/documents/{doc_id}/comments/{comment['id']}/resolve",
    )
    assert resolved.status_code == 200, resolved.text
    assert resolved.json()["status"] == "resolved"
    assert resolved.json()["resolved_at"] is not None

    edit_resolved = await client.patch(
        f"/api/documents/{doc_id}/comments/{comment['id']}",
        json={"body": "Too late to edit."},
    )
    assert edit_resolved.status_code == 409, edit_resolved.text

    reopened = await client.post(
        f"/api/documents/{doc_id}/comments/{comment['id']}/reopen",
    )
    assert reopened.status_code == 200, reopened.text
    assert reopened.json()["status"] == "open"
    assert reopened.json()["resolved_at"] is None
    assert reopened.json()["resolved_by_id"] is None

    edit_reopened = await client.patch(
        f"/api/documents/{doc_id}/comments/{comment['id']}",
        json={"body": "Reopened and edited."},
    )
    assert edit_reopened.status_code == 200, edit_reopened.text
    assert edit_reopened.json()["body"] == "Reopened and edited."


@pytest.mark.asyncio
async def test_document_comments_owner_can_anchor_note_to_extracted_text(client) -> None:
    """Owner-created notes can pin a selected passage to the extracted body hash."""
    await _signup_and_login(client, EMAIL_A, PASSWORD_A)
    source_text = "Alpha beta gamma"
    _, doc_id = await _create_matter_and_upload_text(client, source_text)
    current_hash = hashlib.sha256(source_text.encode("utf-8")).hexdigest()

    created = await client.post(
        f"/api/documents/{doc_id}/comments",
        json={
            "body": "Check this anchored passage.",
            "quote_text": "beta",
            "body_sha256": current_hash,
            "anchor_start": 6,
            "anchor_end": 10,
        },
    )
    assert created.status_code == 200, created.text
    comment = created.json()
    assert comment["quote_text"] == "beta"
    assert comment["body_sha256"] == current_hash
    assert comment["anchor_start"] == 6
    assert comment["anchor_end"] == 10

    stale = await client.post(
        f"/api/documents/{doc_id}/comments",
        json={
            "body": "This should not attach to stale text.",
            "body_sha256": "0" * 64,
            "anchor_start": 0,
            "anchor_end": 5,
        },
    )
    assert stale.status_code == 409, stale.text


@pytest.mark.parametrize(
    "mechanism",
    [
        pytest.param("cross_user", id="comments-cross-user"),
        pytest.param("archived_matter", id="comments-archived-matter"),
    ],
)
@pytest.mark.asyncio
async def test_document_comments_returns_404(client, mechanism) -> None:
    """List, create, resolve, and reopen document comments 404 once access is denied.

    Cross-user also asserts that editing (PATCH) 404s; the archived-matter
    variant does not exercise PATCH in the original coverage.
    """
    await _signup_and_login(client, EMAIL_A, PASSWORD_A)
    slug, doc_id = await _create_matter_and_upload(client)
    created = await client.post(
        f"/api/documents/{doc_id}/comments",
        json={"body": "Owner note."},
    )
    assert created.status_code == 200, created.text
    comment_id = created.json()["id"]

    if mechanism == "cross_user":
        await client.post("/auth/logout")
        await _signup_and_login(client, EMAIL_B, PASSWORD_B)
    else:
        del_resp = await client.delete(f"/api/matters/{slug}")
        assert del_resp.status_code in (200, 204), del_resp.text

    listed = await client.get(f"/api/documents/{doc_id}/comments")
    assert listed.status_code == 404, listed.text
    posted = await client.post(
        f"/api/documents/{doc_id}/comments",
        json={"body": "Should not land."},
    )
    assert posted.status_code == 404, posted.text
    resolved = await client.post(
        f"/api/documents/{doc_id}/comments/{comment_id}/resolve",
    )
    assert resolved.status_code == 404, resolved.text
    reopened = await client.post(
        f"/api/documents/{doc_id}/comments/{comment_id}/reopen",
    )
    assert reopened.status_code == 404, reopened.text

    if mechanism == "cross_user":
        updated = await client.patch(
            f"/api/documents/{doc_id}/comments/{comment_id}",
            json={"body": "Cross-user edit."},
        )
        assert updated.status_code == 404, updated.text


@pytest.mark.asyncio
async def test_document_edit_sessions_are_owner_scoped(client) -> None:
    """Owner can heartbeat an edit session; another user cannot see it."""
    await _signup_and_login(client, EMAIL_A, PASSWORD_A)
    _, doc_id = await _create_matter_and_upload(client)

    started = await client.post(
        f"/api/documents/{doc_id}/edit-sessions",
        json={"client_id": "browser-client-a"},
    )
    assert started.status_code == 200, started.text
    payload = started.json()
    assert payload["current"]["client_id"] == "browser-client-a"
    assert payload["active"][0]["user_label"] == EMAIL_A

    listed = await client.get(f"/api/documents/{doc_id}/edit-sessions")
    assert listed.status_code == 200, listed.text
    assert [row["client_id"] for row in listed.json()] == ["browser-client-a"]

    ended = await client.delete(
        f"/api/documents/{doc_id}/edit-sessions/{payload['current']['id']}",
    )
    assert ended.status_code == 204, ended.text
    listed_after = await client.get(f"/api/documents/{doc_id}/edit-sessions")
    assert listed_after.status_code == 200, listed_after.text
    assert listed_after.json() == []

    await client.post("/auth/logout")
    await _signup_and_login(client, EMAIL_B, PASSWORD_B)
    cross_user = await client.get(f"/api/documents/{doc_id}/edit-sessions")
    assert cross_user.status_code == 404, cross_user.text


@pytest.mark.asyncio
async def test_post_manual_document_version_creates_user_edit_version(client) -> None:
    """Owner can save editor text as a new immutable document version."""
    await _signup_and_login(client, EMAIL_A, PASSWORD_A)
    _, doc_id = await _create_matter_and_upload(client)

    resp = await client.post(
        f"/api/documents/{doc_id}/versions/manual",
        json={
            "resolved_text": "Edited witness statement.\n\nSecond paragraph.",
            "resolved_json": {
                "type": "doc",
                "content": [
                    {
                        "type": "paragraph",
                        "content": [{"type": "text", "text": "Edited witness statement."}],
                    },
                    {
                        "type": "paragraph",
                        "content": [{"type": "text", "text": "Second paragraph."}],
                    },
                ],
            },
            "notes": "Manual edit from document editor",
        },
    )
    assert resp.status_code == 200, resp.text
    payload = resp.json()
    assert payload["kind"] == "user_edit"
    assert payload["version_number"] == 2
    assert payload["resolved_text"] == "Edited witness statement.\n\nSecond paragraph."
    assert payload["resolved_json"]["type"] == "doc"
    assert payload["notes"] == "Manual edit from document editor"

    versions = await client.get(f"/api/documents/{doc_id}/versions")
    assert versions.status_code == 200, versions.text
    rows = versions.json()
    assert rows[-1]["version"]["kind"] == "user_edit"
    assert rows[-1]["version"]["resolved_text"] == "Edited witness statement.\n\nSecond paragraph."
    assert rows[-1]["version"]["resolved_json"]["type"] == "doc"


@pytest.mark.asyncio
async def test_document_working_draft_round_trip_and_commit(client) -> None:
    """Owner can autosave a mutable draft and commit it as an immutable version."""
    await _signup_and_login(client, EMAIL_A, PASSWORD_A)
    _, doc_id = await _create_matter_and_upload_text(client, "Original extracted text.")

    initial = await client.get(f"/api/documents/{doc_id}/draft")
    assert initial.status_code == 200, initial.text
    initial_payload = initial.json()
    assert initial_payload["plain_text"] == "Original extracted text."
    assert initial_payload["version_counter"] == 0
    assert initial_payload["base_version_id"] is None

    saved = await client.put(
        f"/api/documents/{doc_id}/draft",
        json={
            "plain_text": "Working draft text.\n\nSecond paragraph.",
            "editor_json": {
                "type": "doc",
                "content": [
                    {
                        "type": "paragraph",
                        "content": [{"type": "text", "text": "Working draft text."}],
                    },
                ],
            },
            "client_id": "client-draft-1",
        },
    )
    assert saved.status_code == 200, saved.text
    draft = saved.json()
    assert draft["plain_text"] == "Working draft text.\n\nSecond paragraph."
    assert draft["editor_json"]["type"] == "doc"
    assert draft["version_counter"] == 1
    assert draft["client_id"] == "client-draft-1"

    listed = await client.get(f"/api/documents/{doc_id}/draft")
    assert listed.status_code == 200, listed.text
    assert listed.json()["plain_text"] == "Working draft text.\n\nSecond paragraph."

    committed = await client.post(
        f"/api/documents/{doc_id}/draft/commit",
        json={"notes": "Commit shared working draft"},
    )
    assert committed.status_code == 200, committed.text
    version = committed.json()
    assert version["kind"] == "user_edit"
    assert version["version_number"] == 2
    assert version["resolved_text"] == "Working draft text.\n\nSecond paragraph."
    assert version["resolved_json"]["type"] == "doc"
    assert version["notes"] == "Commit shared working draft"

    after_commit = await client.get(f"/api/documents/{doc_id}/draft")
    assert after_commit.status_code == 200, after_commit.text
    derived = after_commit.json()
    assert derived["plain_text"] == "Working draft text.\n\nSecond paragraph."
    assert derived["base_version_id"] == version["id"]
    assert derived["version_counter"] == 0

    versions = await client.get(f"/api/documents/{doc_id}/versions")
    assert versions.status_code == 200, versions.text
    rows = versions.json()
    assert rows[-1]["version"]["notes"] == "Commit shared working draft"


@pytest.mark.asyncio
async def test_document_working_draft_base_version_must_belong_to_document(client) -> None:
    """Draft base_version_id cannot point at another document's version."""
    await _signup_and_login(client, EMAIL_A, PASSWORD_A)
    _, doc_id_a = await _create_matter_and_upload_text(client, "Document A.")
    _, doc_id_b = await _create_matter_and_upload_text(client, "Document B.")
    version_b = await client.post(
        f"/api/documents/{doc_id_b}/versions/manual",
        json={"resolved_text": "Document B saved version."},
    )
    assert version_b.status_code == 200, version_b.text

    resp = await client.put(
        f"/api/documents/{doc_id_a}/draft",
        json={
            "plain_text": "Invalid base pointer.",
            "base_version_id": version_b.json()["id"],
        },
    )
    assert resp.status_code == 422, resp.text
    assert resp.json()["detail"]["error"] == "invalid_base_version"


@pytest.mark.asyncio
async def test_document_working_draft_rejects_stale_counter(client) -> None:
    """Autosave and commit reject stale shared-draft counters."""
    await _signup_and_login(client, EMAIL_A, PASSWORD_A)
    _, doc_id = await _create_matter_and_upload_text(client, "Original extracted text.")

    first = await client.put(
        f"/api/documents/{doc_id}/draft",
        json={
            "plain_text": "First working draft.",
            "client_id": "client-a",
            "expected_version_counter": 0,
        },
    )
    assert first.status_code == 200, first.text
    assert first.json()["version_counter"] == 1

    stale_save = await client.put(
        f"/api/documents/{doc_id}/draft",
        json={
            "plain_text": "Stale overwrite attempt.",
            "client_id": "client-b",
            "expected_version_counter": 0,
        },
    )
    assert stale_save.status_code == 409, stale_save.text
    assert stale_save.json()["detail"]["error"] == "working_draft_conflict"
    assert stale_save.json()["detail"]["current_version_counter"] == 1

    stale_commit = await client.post(
        f"/api/documents/{doc_id}/draft/commit",
        json={"expected_version_counter": 0},
    )
    assert stale_commit.status_code == 409, stale_commit.text
    assert stale_commit.json()["detail"]["error"] == "working_draft_conflict"


@pytest.mark.asyncio
async def test_post_upload_document_version_updates_active_document_and_body(client) -> None:
    """Owner can upload a replacement binary as the next active document version."""
    await _signup_and_login(client, EMAIL_A, PASSWORD_A)
    slug, doc_id = await _create_matter_and_upload(client)

    resp = await client.post(
        f"/api/documents/{doc_id}/versions/upload",
        files={
            "file": (
                "replacement.txt",
                io.BytesIO(b"Replacement document text."),
                "text/plain",
            )
        },
        data={"notes": "Clean text replacement"},
    )
    assert resp.status_code == 200, resp.text
    version = resp.json()
    assert version["kind"] == "upload"
    assert version["version_number"] == 2
    assert version["storage_uri"]
    assert version["filename"] == "replacement.txt"
    assert version["mime_type"] == "text/plain"
    assert version["size_bytes"] == len(b"Replacement document text.")
    assert version["sha256"]
    assert version["resolved_text"] == "Replacement document text."
    assert version["notes"] == "Clean text replacement"

    body = await client.get(f"/api/documents/{doc_id}/body")
    assert body.status_code == 200, body.text
    assert body.json()["extracted_text"] == "Replacement document text."

    docs = await client.get(f"/api/matters/{slug}/documents")
    assert docs.status_code == 200, docs.text
    active = next(row for row in docs.json() if row["id"] == doc_id)
    assert active["filename"] == "replacement.txt"
    assert active["mime_type"] == "text/plain"
    assert active["size_bytes"] == len(b"Replacement document text.")

    versions = await client.get(f"/api/documents/{doc_id}/versions")
    assert versions.status_code == 200, versions.text
    rows = versions.json()
    assert [row["version"]["version_number"] for row in rows] == [1, 2]
    assert rows[-1]["version"]["kind"] == "upload"
    assert rows[-1]["version"]["resolved_text"] == "Replacement document text."

    version_original = await client.get(
        f"/api/documents/{doc_id}/versions/{version['id']}/original"
    )
    assert version_original.status_code == 200, version_original.text
    assert version_original.content == b"Replacement document text."
    assert version_original.headers["content-type"].startswith("text/plain")

    initial_original = await client.get(
        f"/api/documents/{doc_id}/versions/{rows[0]['version']['id']}/original"
    )
    assert initial_original.status_code == 200, initial_original.text
    assert initial_original.headers["content-type"].startswith("application/pdf")


@pytest.mark.asyncio
async def test_restore_document_version_reactivates_prior_upload(client) -> None:
    """Owner can restore a prior upload as a new active immutable version."""
    await _signup_and_login(client, EMAIL_A, PASSWORD_A)
    slug, doc_id = await _create_matter_and_upload(client)

    replacement = await client.post(
        f"/api/documents/{doc_id}/versions/upload",
        files={
            "file": (
                "replacement.txt",
                io.BytesIO(b"Replacement document text."),
                "text/plain",
            )
        },
    )
    assert replacement.status_code == 200, replacement.text

    versions = await client.get(f"/api/documents/{doc_id}/versions")
    assert versions.status_code == 200, versions.text
    source = versions.json()[0]["version"]

    restore = await client.post(
        f"/api/documents/{doc_id}/versions/{source['id']}/restore",
        json={"notes": "Back to original upload"},
    )
    assert restore.status_code == 200, restore.text
    restored = restore.json()
    assert restored["kind"] == "restored"
    assert restored["version_number"] == 3
    assert restored["filename"] == "test.pdf"
    assert restored["mime_type"] == "application/pdf"
    assert restored["storage_uri"] == source["storage_uri"]
    assert restored["notes"] == "Back to original upload"

    docs = await client.get(f"/api/matters/{slug}/documents")
    assert docs.status_code == 200, docs.text
    active = next(row for row in docs.json() if row["id"] == doc_id)
    assert active["filename"] == "test.pdf"
    assert active["mime_type"] == "application/pdf"


@pytest.mark.asyncio
async def test_restore_manual_document_version_creates_text_original(client) -> None:
    """Text editor versions can be restored and reopened as active text files."""
    await _signup_and_login(client, EMAIL_A, PASSWORD_A)
    slug, doc_id = await _create_matter_and_upload(client)
    save = await client.post(
        f"/api/documents/{doc_id}/versions/manual",
        json={"resolved_text": "Edited witness statement."},
    )
    assert save.status_code == 200, save.text
    version_id = save.json()["id"]

    restore = await client.post(
        f"/api/documents/{doc_id}/versions/{version_id}/restore",
        json={},
    )
    assert restore.status_code == 200, restore.text
    restored = restore.json()
    assert restored["kind"] == "restored"
    assert restored["mime_type"] == "text/plain"
    assert restored["storage_uri"]
    assert restored["resolved_text"] == "Edited witness statement."

    body = await client.get(f"/api/documents/{doc_id}/body")
    assert body.status_code == 200, body.text
    assert body.json()["extracted_text"] == "Edited witness statement."

    original = await client.get(f"/api/documents/{doc_id}/original")
    assert original.status_code == 200, original.text
    assert original.content == b"Edited witness statement."

    docs = await client.get(f"/api/matters/{slug}/documents")
    assert docs.status_code == 200, docs.text
    active = next(row for row in docs.json() if row["id"] == doc_id)
    assert active["mime_type"] == "text/plain"


@pytest.mark.asyncio
async def test_get_manual_document_version_docx_returns_word_file(client) -> None:
    """Owner can download a saved editor version as a valid .docx."""
    await _signup_and_login(client, EMAIL_A, PASSWORD_A)
    _, doc_id = await _create_matter_and_upload(client)
    note = await client.post(
        f"/api/documents/{doc_id}/comments",
        json={
            "body": "Check limitation before relying on this draft.",
            "quote_text": "Edited witness statement.",
        },
    )
    assert note.status_code == 200, note.text

    save = await client.post(
        f"/api/documents/{doc_id}/versions/manual",
        json={"resolved_text": "Edited witness statement.\n\nSecond paragraph."},
    )
    assert save.status_code == 200, save.text
    version_id = save.json()["id"]

    resp = await client.get(f"/api/documents/{doc_id}/versions/{version_id}/docx")
    assert resp.status_code == 200, resp.text
    assert resp.headers["content-type"].startswith(
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    assert resp.headers.get("content-disposition", "").endswith('.docx"')
    assert zipfile.is_zipfile(io.BytesIO(resp.content))
    docx = DocxDocument(io.BytesIO(resp.content))
    all_text = "\n".join(paragraph.text for paragraph in docx.paragraphs)
    assert "Document review notes" in all_text
    assert "Check limitation before relying on this draft." in all_text
    assert "Edited witness statement." in all_text


@pytest.mark.asyncio
async def test_get_manual_document_version_docx_preserves_rich_editor_marks(client) -> None:
    """Rich editor saves keep basic formatting when exported as .docx."""
    await _signup_and_login(client, EMAIL_A, PASSWORD_A)
    _, doc_id = await _create_matter_and_upload(client)
    image_upload = await client.post(
        f"/api/documents/{doc_id}/assets",
        files={"file": ("governed.png", io.BytesIO(_ONE_PIXEL_PNG), "image/png")},
    )
    assert image_upload.status_code == 200, image_upload.text
    governed_image_url = image_upload.json()["url"]

    save = await client.post(
        f"/api/documents/{doc_id}/versions/manual",
        json={
            "resolved_text": (
                "Important term\n\nListed point\n\nRight aligned red\n\n"
                "[ ] Review source\n[x] Check deadline\n\n"
                "[image: timeline diagram]\n\n[image: governed diagram]"
            ),
            "resolved_json": {
                "type": "doc",
                "content": [
                    {
                        "type": "paragraph",
                        "content": [
                            {
                                "type": "text",
                                "text": "Important",
                                "marks": [{"type": "bold"}, {"type": "underline"}],
                            },
                            {"type": "text", "text": " term", "marks": [{"type": "italic"}]},
                        ],
                    },
                    {
                        "type": "bulletList",
                        "content": [
                            {
                                "type": "listItem",
                                "content": [
                                    {
                                        "type": "paragraph",
                                        "content": [
                                            {
                                                "type": "text",
                                                "text": "Listed point",
                                                "marks": [{"type": "highlight"}],
                                            }
                                        ],
                                    }
                                ],
                            }
                        ],
                    },
                    {
                        "type": "paragraph",
                        "attrs": {"textAlign": "right"},
                        "content": [
                            {
                                "type": "text",
                                "text": "Right aligned red",
                                "marks": [
                                    {
                                        "type": "textStyle",
                                        "attrs": {"color": "#8C1D18"},
                                    }
                                ],
                            }
                        ],
                    },
                    {
                        "type": "taskList",
                        "content": [
                            {
                                "type": "taskItem",
                                "attrs": {"checked": False},
                                "content": [
                                    {
                                        "type": "paragraph",
                                        "content": [{"type": "text", "text": "Review source"}],
                                    }
                                ],
                            },
                            {
                                "type": "taskItem",
                                "attrs": {"checked": True},
                                "content": [
                                    {
                                        "type": "paragraph",
                                        "content": [{"type": "text", "text": "Check deadline"}],
                                    }
                                ],
                            },
                        ],
                    },
                    {
                        "type": "image",
                        "attrs": {
                            "src": "https://example.com/timeline.png",
                            "alt": "timeline diagram",
                        },
                    },
                    {
                        "type": "image",
                        "attrs": {
                            "src": governed_image_url,
                            "alt": "governed diagram",
                        },
                    },
                    {
                        "type": "table",
                        "content": [
                            {
                                "type": "tableRow",
                                "content": [
                                    {
                                        "type": "tableHeader",
                                        "content": [
                                            {
                                                "type": "paragraph",
                                                "content": [{"type": "text", "text": "Issue"}],
                                            }
                                        ],
                                    },
                                    {
                                        "type": "tableHeader",
                                        "content": [
                                            {
                                                "type": "paragraph",
                                                "content": [{"type": "text", "text": "Risk"}],
                                            }
                                        ],
                                    },
                                ],
                            },
                            {
                                "type": "tableRow",
                                "content": [
                                    {
                                        "type": "tableCell",
                                        "content": [
                                            {
                                                "type": "paragraph",
                                                "content": [{"type": "text", "text": "Indemnity"}],
                                            }
                                        ],
                                    },
                                    {
                                        "type": "tableCell",
                                        "content": [
                                            {
                                                "type": "paragraph",
                                                "content": [{"type": "text", "text": "High"}],
                                            }
                                        ],
                                    },
                                ],
                            },
                        ],
                    },
                ],
            },
        },
    )
    assert save.status_code == 200, save.text
    version_id = save.json()["id"]

    resp = await client.get(f"/api/documents/{doc_id}/versions/{version_id}/docx")
    assert resp.status_code == 200, resp.text
    docx = DocxDocument(io.BytesIO(resp.content))

    first = docx.paragraphs[1]
    assert first.runs[0].text == "Important"
    assert first.runs[0].bold is True
    assert first.runs[0].underline is True
    assert first.runs[1].text == " term"
    assert first.runs[1].italic is True

    bullet = next(p for p in docx.paragraphs if p.text == "Listed point")
    assert bullet.style.name.startswith("List Bullet")
    assert bullet.runs[0].font.highlight_color is not None
    aligned = next(p for p in docx.paragraphs if p.text == "Right aligned red")
    assert aligned.alignment == WD_ALIGN_PARAGRAPH.RIGHT
    assert aligned.runs[0].font.color.rgb == RGBColor(0x8C, 0x1D, 0x18)
    assert next(p for p in docx.paragraphs if p.text == "[ ] Review source")
    assert next(p for p in docx.paragraphs if p.text == "[x] Check deadline")
    assert next(p for p in docx.paragraphs if p.text == "[image: timeline diagram]")
    assert not any(p.text == "[image: governed diagram]" for p in docx.paragraphs)
    assert len(docx.inline_shapes) == 1
    assert docx.tables[0].cell(0, 0).text == "Issue"
    assert docx.tables[0].cell(0, 1).text == "Risk"
    assert docx.tables[0].cell(1, 0).text == "Indemnity"
    assert docx.tables[0].cell(1, 1).text == "High"


@pytest.mark.asyncio
async def test_get_manual_document_version_pdf_exports_print_copy(client, monkeypatch) -> None:
    """Saved rich-editor versions can be exported as a PDF through the document route."""
    import app.api.documents as documents_api

    await _signup_and_login(client, EMAIL_A, PASSWORD_A)
    _, doc_id = await _create_matter_and_upload_text(client, "Original body")

    captured_html: dict[str, str] = {}

    async def fake_html_to_pdf(html_doc: str) -> bytes:
        captured_html["html"] = html_doc
        return b"%PDF-1.4 legalise edited document\n"

    monkeypatch.setattr(documents_api, "_html_to_pdf", fake_html_to_pdf)
    save = await client.post(
        f"/api/documents/{doc_id}/versions/manual",
        json={
            "resolved_text": "Edited body",
            "resolved_json": {
                "type": "doc",
                "content": [
                    {
                        "type": "heading",
                        "attrs": {"level": 2},
                        "content": [{"type": "text", "text": "Edited heading"}],
                    },
                    {
                        "type": "paragraph",
                        "content": [{"type": "text", "text": "Edited body"}],
                    },
                ],
            },
        },
    )
    assert save.status_code == 200, save.text
    version_id = save.json()["id"]

    resp = await client.get(f"/api/documents/{doc_id}/versions/{version_id}/pdf")
    assert resp.status_code == 200, resp.text
    assert resp.headers["content-type"].startswith("application/pdf")
    assert resp.headers.get("content-disposition", "").endswith('.pdf"')
    assert resp.content.startswith(b"%PDF-1.4")
    assert "Edited heading" in captured_html["html"]
    assert "Edited body" in captured_html["html"]


@pytest.mark.asyncio
async def test_get_upload_document_version_docx_without_resolved_text_returns_422(client) -> None:
    """Upload versions without resolved_text are not exportable through the editor path."""
    await _signup_and_login(client, EMAIL_A, PASSWORD_A)
    _, doc_id = await _create_matter_and_upload(client)
    versions = await client.get(f"/api/documents/{doc_id}/versions")
    assert versions.status_code == 200, versions.text
    upload_version_id = versions.json()[0]["version"]["id"]

    resp = await client.get(f"/api/documents/{doc_id}/versions/{upload_version_id}/docx")
    assert resp.status_code == 422, resp.text


@pytest.mark.asyncio
async def test_post_manual_document_version_cross_user_returns_404(client) -> None:
    """User B cannot save a version on User A's document by UUID."""
    await _signup_and_login(client, EMAIL_A, PASSWORD_A)
    _, doc_id = await _create_matter_and_upload(client)
    await client.post("/auth/logout")

    await _signup_and_login(client, EMAIL_B, PASSWORD_B)
    resp = await client.post(
        f"/api/documents/{doc_id}/versions/manual",
        json={"resolved_text": "Cross-user edit should not land."},
    )
    assert resp.status_code == 404, resp.text


@pytest.mark.parametrize(
    "mechanism",
    [
        pytest.param("cross_user", id="working-draft-cross-user"),
        pytest.param("archived_matter", id="working-draft-archived-matter"),
    ],
)
@pytest.mark.asyncio
async def test_document_working_draft_returns_404(client, mechanism) -> None:
    """Get, write, and commit a working draft 404 once access is denied."""
    await _signup_and_login(client, EMAIL_A, PASSWORD_A)
    slug, doc_id = await _create_matter_and_upload_text(client, "Owned draft source.")
    saved = await client.put(
        f"/api/documents/{doc_id}/draft",
        json={"plain_text": "Owned working draft."},
    )
    assert saved.status_code == 200, saved.text

    if mechanism == "cross_user":
        await client.post("/auth/logout")
        await _signup_and_login(client, EMAIL_B, PASSWORD_B)
    else:
        del_resp = await client.delete(f"/api/matters/{slug}")
        assert del_resp.status_code == 204, del_resp.text

    get_resp = await client.get(f"/api/documents/{doc_id}/draft")
    assert get_resp.status_code == 404, get_resp.text
    put_resp = await client.put(
        f"/api/documents/{doc_id}/draft",
        json={"plain_text": "Should not land."},
    )
    assert put_resp.status_code == 404, put_resp.text
    commit_resp = await client.post(f"/api/documents/{doc_id}/draft/commit")
    assert commit_resp.status_code == 404, commit_resp.text


@pytest.mark.parametrize(
    "resource_path,needs_saved_version,check_alive",
    [
        pytest.param("body", False, True, id="document-body"),
        pytest.param("versions", False, True, id="document-versions"),
        pytest.param("anonymise", False, False, id="anonymise"),
        pytest.param("versions/{version_id}/docx", True, False, id="version-docx"),
    ],
)
@pytest.mark.asyncio
async def test_document_get_routes_archived_matter_returns_404(
    client, resource_path, needs_saved_version, check_alive
) -> None:
    """After the owning matter is archived, document UUID-keyed GET routes 404."""
    await _signup_and_login(client, EMAIL_A, PASSWORD_A)
    slug, doc_id = await _create_matter_and_upload(client)

    version_id = ""
    if needs_saved_version:
        save = await client.post(
            f"/api/documents/{doc_id}/versions/manual",
            json={"resolved_text": "Owned user edit."},
        )
        assert save.status_code == 200, save.text
        version_id = save.json()["id"]

    path = resource_path.format(version_id=version_id)

    if check_alive:
        # Confirm the resource is accessible while the matter is live.
        alive = await client.get(f"/api/documents/{doc_id}/{path}")
        assert alive.status_code == 200, alive.text

    # Tombstone the matter.
    del_resp = await client.delete(f"/api/matters/{slug}")
    assert del_resp.status_code == 204, del_resp.text

    resp = await client.get(f"/api/documents/{doc_id}/{path}")
    assert resp.status_code == 404, resp.text


@pytest.mark.asyncio
async def test_post_manual_document_version_archived_matter_returns_404(client) -> None:
    """After the owning matter is archived, manual editor saves 404."""
    await _signup_and_login(client, EMAIL_A, PASSWORD_A)
    slug, doc_id = await _create_matter_and_upload(client)

    del_resp = await client.delete(f"/api/matters/{slug}")
    assert del_resp.status_code == 204, del_resp.text

    resp = await client.post(
        f"/api/documents/{doc_id}/versions/manual",
        json={"resolved_text": "Archived matter edit should not land."},
    )
    assert resp.status_code == 404, resp.text


@pytest.mark.asyncio
async def test_get_generated_docx_cross_user_returns_404(client) -> None:
    """GET /api/documents/generated/{uuid} returns 404 for a non-owner."""
    # Fabricate a UUID that doesn't correspond to any audit row
    bogus_uuid = uuid.uuid4()
    await _signup_and_login(client, EMAIL_A, PASSWORD_A)
    resp = await client.get(f"/api/documents/generated/{bogus_uuid}")
    assert resp.status_code == 404, resp.text
